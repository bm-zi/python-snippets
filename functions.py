#!/usr/bin/python3
# -*- coding: utf-8 -*-

# Author: bm-zi

def f1():
    # Chapter 1 - Python Basics
    print('''Chapter 1 - Python Basics
    . Entering Expressions into the Interactive Shell
    . The Integer, Floating-Point, and String Data Types
    . String Concatenation and Replication
    . Storing Values in Variables
    . Your First Program
        ''')

def f2():
    # Math Operators
    print(2+2)  # 4
    print(2)    # 2
    print(2 + 3 * 6)    # 20
    print((2 + 3) * 6)  # 30
    print(48565878 * 578453)    # 28093077826734
    print(2 ** 8)   # 256
    print(23 / 7)   # 3.2857142857142856
    print(23 // 7)  # 3
    print(23 % 7)   # 2
    print(2    +    2)  # 4
    print((5 - 1) * ((7 + 1) / (3 - 1)))  # 16.0
    print()  # blank line in python3

def f3():
    # String concatenation and replication
    print('Alice' + 'Bob') # AliceBob
    print('Alice' * 5)  # AliceAliceAliceAliceAlice

def f4():
    # First Python Program - print user input

    print('Hello world!')
    print('What is your name?') # ask for their name,
                                # name has to be entered as a str-
                                # ing like 'Jim' in python2
    myName = input()
    print('It is good to meet you, ' + myName)
    print('The length of your name is:')
    print(len(myName))
    print('What is your age?') # ask for their age
    myAge = input()
    print('You will be ' + str(int(myAge) + 1) + ' in a year.')

def f5():
    # Function len() - string length
    print(len('hello'))
    print(len('My very energetic monster just scarfed nachos.'))

def f6():
    # Function str()
    print('I am ' + str(29) + ' years old.')

    # you get an error if you just use 29 without function str()

def f7():
    # Function str(), int(), float
    # The str() , int() , and float() functions will evaluate to the string,
    # integer, and floating-point forms of the value you pass, respectively.
    print(str(0))     # 0
    print(str(-3.14)) # -3.14
    print(int('42'))  # 42
    print(int('-99')) # -99
    print(int(1.25))  # 1
    print(int(1.99))  # 1
    print(float('3.14'))  # 3.14
    print(float(10))  # 10.0

def f8():
    # input() function
    # input() function always returns a string,
    # to change it to an integer use int() function
    spam = input()
    spam = int(spam)
    print(spam)   # If you type a string that cannot be converted to integer,
                    # you will get an error shows that the variable is not defined
                    # These are error:  int('99.99') int('twelve')

    print(int(7.7))      # 7
    print(int(7.7) + 1)  # 8
    print(42 == '42')    # False
    print(42 == 42.0)    # True
    print(42.0 == 0042.000)   # True
    print()
    print()
    print()

def f9():
    print('''Chapter 2 - Flow Control:
    . Boolean Values
    . Comparison Operators
    . Boolean Operators
    . Mixing Boolean and Comparison Operators
    . Elements of Flow Control (if, else, ilif, while, break, continue, for and rang)
    . Flow Control Statements
    . Importing Modules
    . Ending a Program Early with sys.exit()
    ''')

def f10():
    # Boolean data types
    # --- CHAPTER 2 ---
    # Boolean data types are having values: False and True

    print(42 == 42)   # True
    print(42 == 99)   # False
    print(2 != 3)     # True
    print(2 != 2)     # False

    # The == and != operators can actually work with values of any data type.

    print('hello' == 'hello')   # True
    print('hello' == 'Hello')   # False
    print('dog' != 'cat')       # True
    print(True == True)         # True
    print(True != False)        # True
    print(42 == 42.0)           # True
    print(42 == '42')           # False

    print(42 < 100)         # True
    print(42 > 100)         # False
    print(42 < 42)          # False

    eggCount = 42
    print(eggCount <= 42)   # True

    myAge = 29
    print(myAge >= 10)      # True

def f11():
    # Boolean Operators
    # and operator
    print(True and True)    # True
    print(True and False)   # False
    print(False and True)   # False
    print(False and False)  # False


    # or operator
    print(True or True)     # True
    print(True or False)    # True
    print(False or True)    # True
    print(False or False)   # False

    # not operator
    print(not True)     # False
    print(not False)    # True

def f12():
    # Mixing Boolean and Comparison Operators
    print((4 < 5) and (5 < 6))      # True
    print((4 < 5) and (9 < 6))      # False
    print((1 == 2) or (2 == 2))     # True
    print(2 + 2 == 4 and not 2 + 2 == 5 and 2 * 2 == 2 + 2)     # True

def f13():
    # if,elif,else conditional
    name='alice'
    age=20

    if name == 'Alice':
        print('Hi, Alice.')
    elif age < 12:
        print('You are not Alice, kiddo.')
    else:
        print('You are neither Alice nor a little kid.')


    spam = 0
    if spam < 5:
        print('Hello, world.')
        spam = spam + 1

def f14():
    # while loop
    spam = 0
    while spam < 5:
        print('Hello, world.')
        spam = spam + 1

def f15():
    # An Annoying while Loop
    name = 'Al'
    while name != 'your name':
        print('Please type your name.')
        name = input()
    print('Thank you!')

def f16():
    # break Statements, breaks the loop
    while True:
        print('Please type your name.')
        name = input()
        if name == 'your name':
            break
    print('Thank you!')

def f17():
    # continue statement, keeps running the loop start from begining
    while True:
        print('Who are you?')
        name = input()
        if name != 'Joe':
            continue   # restarts the loop
        print('Hello, Joe. What is the password? (It is a fish.)')
        password = input()
        if password == 'swordfish':
            break      # stops the loop
    print('Access granted.')

    # Note:
    # A combination of 'if not statement' with 'continue' helps if
    # the condition is not true then loops resets, otherwise it k-
    # eeps going to the next statement in loop.

def f18():
    # empty value is considered False
    name = ''
    while not name:
        print('Entenning the loop start from beginingr your name:')
        name = input()
    print('How many guests will you have?')
    numOfGuests = int(input())
    if numOfGuests:     # if no of guest not empty
        print('Be sure to have enough room for all your guests.')
    print('Done')

def f19():
    # for Loops and the range() Function
    print('My name is')
    for i in range(5):
        print('Jimmy Five Times (' + str(i) + ')')

def f20():
    # while loop equivalent to for loop
    print('My name is')
    i = 0
    while i < 5:
        print('Jimmy Five Times (' + str(i) + ')')
        i = i + 1

def f21():
    # for loop example - sum of numbers from 1 to 100
    total = 0
    for num in range(101):
        total = total + num
    print(total)

def f22():
    # range() function examples
    for i in range(12, 16):
        print(i)


    print('.' * 10)

    for i in range(0, 10, 2):
        print(i)

    print('.' * 10)

    for i in range(5, -1, -1):
        print(i)

def f23():
    # import modules
    # functions like print(), input() , and len() are builtin fuctions;
    # No need to import them.
    # function like randint() are from standard library modules that
    # has to be imported

    # if use "from random import *" ,
    # then no need to to call a function with the random. prefix.

    import random
    for i in range(5):
        print(random.randint(1, 10))

def f24():
    # sys.exit()
    # Ending a Program Early with sys.exit()

    import sys
    while True:
        print('Type exit to exit.')
        response = input()
        if response == 'exit':
            sys.exit()
        print('You typed ' + response + '.')

def f25():
    #
    print('''Chapter 3 - Functions:
    . def Statements with Parameters
    . Return Values and return Statements
    . The None Value
    . Keyword Arguments and print()
    . Local and Global Scope
    . The global Statement
    . Exception Handling''')

def f26():
    # FUNCTIONS
    def hello():

        print('Howdy!')
        print('Howdy!!!')
        print('Hello there.')
    hello()
    hello()
    hello()

def f27():
    # Function with argument
    def hello(name):

        print('Hello ' + name)
    hello('Alice')
    hello('Bob')

def f28():
    # return statement
    import random
    def getAnswer(answerNumber):

        if answerNumber == 1:
            return 'It is certain'
        elif answerNumber == 2:
            return 'It is decidedly so'
        elif answerNumber == 3:
            return 'Yes'
        elif answerNumber == 4:
            return 'Reply hazy try again'
        elif answerNumber == 5:
            return 'Ask again later'
        elif answerNumber == 6:
            return 'Concentrate and ask again'
        elif answerNumber == 7:
            return 'My reply is no'
        elif answerNumber == 8:
            return 'Outlook not so good'
        elif answerNumber == 9:
            return 'Very doubtful'

    r = random.randint(1, 9)
    fortune = getAnswer(r)
    print(fortune)
    # return exits the fuction block

def f29():
    # None type
    # 'None' is a value-without-value
    # (like null or undefined in other languages.)
    spam = print('Hello!') # print function returns no value that
                            # is known as 'None'
    print(None == spam)

    # return without value returns value 'None'
    # continue in for or while loop also returns value 'None'

def f30():
    # Optional keywords for functions
    # functions have optional keyword like end and sep in print() function

    print('Hello', end='')    # The 'end' keyword cause to ignore the new line.
    print('World')            # By defualt print prints a newline after printing
                                # its argument
    print('cats', 'dogs', 'mice')
    print('cats', 'dogs', 'mice', sep=',') # the function print will automatically
                                            # separates multiple string values
                                            # with a single space.

def f31():
    # Local and Global Scopes

    # 1.
    # Local Variables Cannot Be Used in the Global Scope
    # The following snippet generates error if uncommented:
    # def spam():
    #     eggs = 31337
    # spam()
    # print(eggs)

    # 2.
    # Local Scopes Cannot Use Variables in Other Local Scopes
    # The following snippet prints '99' if is uncommented
    def spam1():
        eggs = 99
        bacon1()
        print(eggs)

    def bacon1():
        ham = 101
        eggs = 0

    spam1()
    print('.' * 10)

    # 3.
    # Global Variables Can Be Read from a Local Scope
    # '42' will be printed, for the following snippet
    def spam2():
        # eggs = 21
        print(eggs)    # because no local eggs exists, python con-
                       # siders it a reference to the global vari-
                       # able eggs(42). If a local variale eggs -
                       # exists, it uses that local variable inst-
                       # ead of the variable eggs in global scope.

    eggs = 42
    spam2()
    print(eggs)  # prints 21, if 'eggs = 21' is uncommented.
    print('.' * 10)

    # 4.
    # Local and Global Variables with the Same Name
    # Avoid using the same variable name in different scopes.
    def spam3():
        eggs = 'spam local'
        print(eggs)            # prints 'spam local'

    def bacon3():
        eggs = 'bacon local'
        print(eggs)            # prints 'bacon local'
        spam3()                 # prints 'spam local'
        print(eggs)            # prints 'bacon local'

    eggs = 'global'
    bacon3()
    print(eggs)                # prints 'global'

def f32():
    # Local and Global Variables

    # If you need to modify a global variable from within a funct-
    # ion, use the global statement.
    def spam4():
        global eggs
        eggs = 'spam'
        print(eggs)

    eggs = 'global'
    spam4()        # prints spam
    print(eggs)    # prints 'global'
    print('.' * 10)

    # Rules:
    # 1.   If a variable is being used in the global scope (that -
    #      is, outside of all functions), then it is always a glo-
    #      bal variable.

    # 2.   If there is a global statement for that variable in a -
    #      function, it is a global variable.

    # 3.   Otherwise, if the variable is used in an assignment st-
    #      atement in the function, it is a local variable.

    # 4.   But if the variable is not used in an assignment state-
    #      ment, it is a global variable.

    # To get a better feel for the above rule, here is another sn-
    # ippet:
    def spam5():
        global eggs
        eggs = 'spam'  # this is the global
        print(eggs)

    def bacon5():
        eggs = 'bacon' # this is a local
        print(eggs)

    def ham5():
        print(eggs)    # this is the global

    eggs = 42          # this is the global
    spam5()
    print(eggs)
    bacon5()
    ham5()

    print('.' * 10)
    # If try to use a local variable in a function before you ass-
    # ign a value to it, Python will give you an error.
    def spam6():
        # print(eggs) # if uncommented you get ERROR!
        eggs = 'spam local'
        print(eggs)

    eggs = 'global'
    spam6()

def f33():
    # Exception Handling

    def spam(divideBy):
        try:                        # Potential error code is in try clause
            return 42 / divideBy
        except ZeroDivisionError:   # When error is caught in try clause, then it moves to except block,
                                    # and then gets back to try block and continue the next statement
                                    # after the statement with error
            print('Error: Invalid argument.')

    print(spam(2))
    print(spam(12))
    print(spam(0))
    print(spam(1))

def f34():
    # Out Of Function Exception Handling
    # Difference between the folloing code and the previous one:
    # In previous code any errors that occur in function calls in
    # a try block will also be caught.

    def spam(divideBy):
        return 42 / divideBy

    try:
        print(spam(2))
        print(spam(12))
        print(spam(0))
        print(spam(1))
    except ZeroDivisionError:
        print('Error: Invalid argument.')

def f35():
    # Guessing Game
    import random
    secretNumber = random.randint(1, 20)
    print('I am thinking of a number between 1 and 20.')

    # Ask the player to guess 6 times.
    for guessesTaken in range(1, 7):
        print('Take a guess.')
        guess = int(input())

        if guess < secretNumber:
            print('Your guess is too low.')
        elif guess > secretNumber:
            print('Your guess is too high.')
        else:
            break

    # This condition is the correct guess!
    if guess == secretNumber:
        print('Good job! You guessed my number in ' + str(guessesTaken) + ' guesses!')
    else:
        print('Nope. The number I was thinking of was ' + str(secretNumber))

def f36():
    #
    print('''Chapter 4 - Lists:
    . The List Data Type
    . Working with Lists
    . Augmented Assignment Operators
    . Finding a Value in a List with the index() Method
    . Adding Values to Lists with the append() and insert() Methods
    . Removing Values from Lists with remove()
    . Sorting the Values in a List with the sort() Method
    . List-like Types: Strings and Tuples
    . Mutable and Immutable Data Types
    . The Tuple Data Type
    . Converting Types with the list() and tuple() Functions''')

def f37():
    # Lists - examples
    print([1, 2, 3])
    print(['cat', 'bat', 'rat', 'elephant'])
    print(['hello', 3.1415, True, None, 42])

    # Getting Individual Values in a List with Indexes
    spam = ['cat', 'bat', 'rat', 'elephant']
    print(spam)
    print(spam[0])
    print(spam[1])
    print(spam[2])
    print(spam[3])
    print(['cat', 'bat', 'rat', 'elephant'][3])
    print('Hello ' + spam[0])
    print('The ' + spam[1] + ' ate the ' + spam[0] + '.')
    print(spam[int(1.0)])

def f38():
    # Lists : list of lists
    spam = [['cat', 'bat'], [10, 20, 30, 40, 50]]   # List of lists
    print(spam)
    print(spam[0][1])
    print(spam[1][4])

def f39():
    # Lists : negative indexes
    spam = ['cat', 'bat', 'rat', 'elephant']
    print(spam[-1])     # elephant
    print(spam[-3])     # bat
    print(spam[0:-1])   # ['cat', 'bat', 'rat']
    print('The ' + spam[-1] + ' is afraid of the ' + spam[-3] + '.')

def f40():
    # Lists : getting sublists with slices
    spam = ['cat', 'bat', 'rat', 'elephant']
    print(spam[0:4])    # ['cat', 'bat', 'rat', 'elephant']
    print(spam[1:3])    # ['bat', 'rat']
    print(spam[0:-1])   # ['cat', 'bat', 'rat']
    print(spam[:2])     # ['cat', 'bat']
    print(spam[1:])     # ['bat', 'rat', 'elephant']
    print(spam[:])      # ['cat', 'bat', 'rat', 'elephant']

def f41():
    # Lists : getting a list’s length with len()
    spam = ['cat', 'dog', 'moose']
    print(len(spam))  # 3

def f42():
    # Lists : changing values in a list with indexes
    spam = ['cat', 'bat', 'rat', 'elephant']
    spam[1] = 'aardvark'
    print(spam)
    spam[2] = spam[1]
    print(spam)
    spam[-1] = 12345
    print(spam)

def f43():
    # Lists: list concatenation and list replication

    # It is like strings concatenation and replication
    x = [1, 2, 3] + ['A', 'B', 'C']
    print(x)
    x = ['X', 'Y', 'Z'] * 3
    print(x)
    spam = [1, 2, 3]
    spam = spam + ['A', 'B', 'C']
    print(spam)

def f44():
    # Lists : removing values from lists with del statements
    spam = ['cat', 'bat', 'rat', 'elephant']
    del spam[2]
    print(spam)
    del spam[2]
    print(spam)

def f45():
    # Lists : working with lists

    # instead of using many varibales use lists for data entry
    catNames = []
    while True:
        print('Enter the name of cat ' + str(len(catNames) + 1) + ' (Or enter nothing to stop.):')
        name = input()
        if name == '':
            break
        catNames = catNames + [name] # list concatenation
    print('The cat names are:')
    for name in catNames:
        print('    ' + name)

def f46():
    # Lists : using for loops

    print('simple for loop : ')
    for i in range(4):
        print(i)

    print('using for loop with list : ')
    for i in [0, 1, 2, 3]:
        print(i)

    print('technique of using range(len(someList)) with a for loop : ')
    supplies = ['pens', 'staplers', 'flame-throwers', 'binders']
    for i in range(len(supplies)):
        print('Index ' + str(i) + ' in supplies is: ' + supplies[i])

def f47():
    # Lists : in and not in operators

    # determine whether a value is or isn’t in a list
    print('howdy' in ['hello', 'hi', 'howdy', 'heyas'])
    spam = ['hello', 'hi', 'howdy', 'heyas']
    print(spam)
    print('cat' in spam)
    print('howdy' not in spam)
    print('cat' not in spam)

    print ('...')
    print('Sample program to check if value is in list : ')
    myPets = ['Zophie', 'Pooka', 'Fat-tail']
    print(myPets)
    print('Enter a pet name:')
    name = input()
    if name not in myPets:
        print('I do not have a pet named ' + name)
    else:
        print(name + ' is my pet.')

def f48():
    # Lists : multiple assignment trick
    cat = ['fat', 'black', 'loud']
    size = cat[0]
    color = cat[1]
    disposition = cat[2]
    print(cat)

    print('...  ...  ...  ... ')
    cat = ['fat', 'black', 'loud']
    size, color, disposition = cat
    print(cat)

def f49():
    # Augmented assignment operators

    '''
    statement          Equivalent
    spam = spam + 1    spam += 1
    spam = spam - 1    spam -= 1
    spam = spam * 1    spam *= 1
    spam = spam / 1    spam /= 1
    spam = spam % 1    spam %= 1
    '''

    spam = 42
    spam = spam + 1
    print(spam)

    spam = 42
    spam += 1
    print(spam)

    spam = 'Hello'
    spam += ' world!'
    print(spam)

    bacon = ['Zophie']
    bacon *= 3
    print(bacon)

def f50():
    # Lists : finding a Value in a list with the index() method

    # Methods are the same as functoions, except that is called on value
    # Each data type has its own set of methods.
    # list data type, for ­example, has several useful methods for finding,
    # adding, removing, and manipulating values in a list.

    spam = ['hello', 'hi', 'howdy', 'heyas']
    print(spam.index('hello'))
    print(spam.index('heyas'))

    # When there are duplicates of the value in the list,
    # the index of its first appearance is returned.

    spam = ['Zophie', 'Pooka', 'Fat-tail', 'Pooka']
    print(spam.index('Pooka'))

def f51():
    # Methods - adding values to lists with append() and insert() methods

    spam = ['cat', 'dog', 'bat']
    spam.append('moose')
    print(spam)

    spam = ['cat', 'dog', 'bat']
    spam.insert(1, 'chicken')
    print(spam)

def f52():
    # Lists : removing values from lists with remove()
    spam = ['cat', 'bat', 'rat', 'elephant']
    spam.remove('bat')
    print(spam)

    # If the value appears multiple times in the list,
    # only the first instance of the value will be removed

    spam = ['cat', 'bat', 'rat', 'cat', 'hat', 'cat']
    spam.remove('cat')
    print(spam)

def f53():
    # List : sorting the values in a list with the sort() method
    spam = [2, 5, 3.14, 1, -7]
    spam.sort()
    print(spam)

    spam = ['ants', 'cats', 'dogs', 'badgers', 'elephants']
    spam.sort()
    print(spam)

    # You can also pass True for the reverse keyword argument to have sort()
    # sort the values in reverse order
    spam.sort(reverse=True)
    print(spam)

    # note (1) :
    # You cannot write: spam = spam.sort() ,
    # sort() method sorts the list in place
    # note (2) :
    # cannot sort lists that have both number and string
    # note (3) :
    # sort() uses “ASCIIbetical order” rather than actual alphabetical

    spam = ['Alice', 'ants', 'Bob', 'badgers', 'Carol', 'cats']
    spam.sort()
    print(spam)

    # To sort in alphabetic order:
    spam = ['a', 'z', 'A', 'Z']
    spam.sort(key=str.lower)
    print(spam)

def f54():
    # magic 8 ball with a list

    # much more elegant version of previous 8 ball
    import random
    messages = ['It is certain',
                'It is decidedly so',
    'Yes definitely',
    'Reply hazy try again',
    'Ask again later',
    'Concentrate and ask again',
    'My reply is no',
    'Outlook not so good',
    'Very doubtful']
    print(messages[random.randint(0, len(messages) - 1)])

def f55():
    # list-like types: strings and tuples

    '''
    You can consider a string to be a "list" of single text characters.
    Then you can use indexing; slicing; and using them with for loops,
    with len() , and with the in and not in operators.
    '''
    name = 'Zophie'
    print(name[0])

    print(name[-2])
    print(name[0:4])
    print('Zo' in name)
    print('z' in name)
    print('p' not in name)

    for i in name:
        print('* * * ' + i + ' * * *')

def f56():
    # mutable and immutable data types

    # Lists can have values added, removed, or changed. (mutable)
    # Strings are not. (immutable)

    name = 'Zophie a cat'
    # name[7] = 'the'   # this generates error
    newName = name[0:7] + 'the' + name[8:12]
    print(name)
    print(newName)

    eggs = [1, 2, 3]
    eggs = [4, 5, 6] # The list value in eggs isn’t being changed here; is overwriiten.
    print(eggs)

    print('The actula modification is as following: ')
    eggs = [1, 2, 3]
    del eggs[2]
    del eggs[1]
    del eggs[0]
    eggs.append(4)
    eggs.append(5)
    eggs.append(6)
    print(eggs)

def f57():
    # tuple data type

    # tuples are like lists but use () instead aof []
    # tuples are like strings, cannot have their values modified, appended, or removed.

    eggs = ('hello', 42, 0.5)
    # eggs[1] = 99   # this generates error

    # If you have only one value in your tuple, you can indicate this
    # by placing a trailing comma
    print(type(('hello',)))
    print(type(('hello')))

    '''tuples advantages:
    . When you don't want the values to be changed.(unlike the list)
    . They are faster than lists
    '''
    print('Converting Types with the list() and tuple() Functions')
    print(tuple(['cat', 'dog', 5]))
    print(list(('cat', 'dog', 5)))
    print(list('hello'))

def f58():
    # references

    # assignment for strings and integers value
    spam = 42
    cheese = spam
    spam = 100
    print('spam has different value than cheese: ')
    print(spam)
    print(cheese) # cheese didn't chage

    # assignment foe lists is different, actually we are assiging a
    # list reference to the variable

    spam = [0, 1, 2, 3, 4, 5]
    cheese = spam
    cheese[1] = 'Hello!'
    print('spam and cheese have the same value: ')
    print(spam)
    print(cheese)

    '''
    When you create the list, you assign a reference to it in the spam
    variable. But the next line copies only the list reference in spam
    to cheese , not the list value itself.
    list variables don’t actually contain lists—they contain references
    to lists.
    '''

def f59():
    # Passing References

    '''
    When a function is called, the values of the arguments are copied
    to the parameter variables. For lists and dictionaries, this means
    a copy of the reference is used for the parameter.
    '''

    def eggs(someParameter):
        someParameter.append('Hello')

    spam = [1, 2, 3]
    eggs(spam)
    print(spam)

    '''
    Notice that when eggs() is called, a return value is not used to
    assign a new value to spam. Instead, it modifies the list in place, directly.

    Even though spam and someParameter contain separate references, they
    both refer to the same list. This is why the append('Hello') method call
    inside the function affects the list even after the function call has returned.
    '''

def f60():
    # The copy Module’s copy() and deepcopy() Functions

    '''
    copy.copy() , makes a duplicate copy of a mutable value like a list or dictionary,
    not just a copy of a reference.
    '''

    import copy
    spam = ['A', 'B', 'C', 'D']
    cheese = copy.copy(spam)  # creates a second list that can be
                              # modified ­independently of the first.
    cheese[1] = 42
    print('copy.copy')
    print(spam)
    print(cheese)

    milk = copy.deepcopy(cheese)
    milk[1] = 52
    print('copy.deepcopy')
    print(spam)
    print(cheese)
    print(milk)

def f61():
    grid = [['.','.','.','.','.','.'],
            ['.','0','0','.','.','.'],
            ['0','0','0','0','.','.'],
            ['0','0','0','0','0','.'],
            ['.','0','0','0','0','0'],
            ['0','0','0','0','0','.'],
            ['0','0','0','0','.','.'],
            ['.','0','0','.','.','.'],
            ['.','.','.','.','.','.']]

    print(grid)

def f62():
    # Chapter 5 - Dictionaries and Structuring Data
    print('''Chapter 5 - Dictionaries and Structuring Data:
    . The Dictionary Data Type
        . Using Data Structures to Model Real-World Things
        . The keys(), values(), and items() Methods
        . Checking Whether a Key or Value Exists in a Dictionary
        . The get() Method
        . The setdefault() Method
    . Pretty Printing
    . Using Data Structures to Model Real-World Things
        .   A Tic-Tac-Toe Board
        .   Nested Dictionaries and Lists
    ''')

def f63():
    # Dictionary Data Type
    myCat = {'size': 'fat', 'color': 'grey', 'disposition': "loud" }
    print(myCat['size'])
    print('My cat has ' + myCat['color'] + ' fur.')

    # Dictionaries can use integer values as keys, like lists.
    # There is no item order in dictionaries.
    spam = {12345: 'Luggage Combination', 42: 'The Answer'}
    print(spam)

    print('Dictionaries vs. Lists')
    spam = ['cats', 'dogs', 'moose']
    bacon = ['dogs', 'moose', 'cats']
    print(spam == bacon)

    eggs = {'name': 'Zophie', 'species': 'cat', 'age': '8'}
    ham = {'species': 'cat', 'age': '8', 'name': 'Zophie'}
    print(eggs == ham)

def f64():
    # program to store data
    birthdays = {'Alice': 'Apr 1', 'Bob': 'Dec 12', 'Carol': 'Mar 4'}
    while True:
        print('Enter a name: (blank to quit)')
        name = input()
        if name == '':
            break
        if name in birthdays:
            print(birthdays[name] + ' is the birthday of ' + name)
        else:
            print('I do not have birthday information for ' + name)
            print('What is their birthday?')
            bday = input()
            birthdays[name] = bday
            print('Birthday database updated.')
            print(birthdays)

def f65():
    # Dictionary's methods: keys, values, items
    spam = {'color': 'red', 'age': 42}
    print('Values:')
    for v in spam.values():
        print(v)

    print('Keys::')
    for k in spam.keys():
        print(k)

    print('Items:')
    for i in spam.items():
        print(i)

def f66():
    # Build a list from dictionary's keys
    spam = {'color': 'red', 'age': 42}

    print('spam.keys() returns \'dict_key\' data type')
    print(spam.keys())
    print(list(spam.keys()))

    #
    # The list(spam.keys()) line takes the dict_keys value
    # returned from keys() and passes it to list() , which
    # then returns a list value of ['color', 'age']

def f67():
    # Show dictionary's keys & values by for loop
    spam = {'color': 'red', 'age': 42}
    for k, v in spam.items():
        print('Key: ' + k + ' Value: ' + str(v))

    # spam.items() returns a dict_items data type
    print(spam.items())

def f68():
    # Check if key or value exists in a dictionary,
    # by using in and not operators

    spam = {'name': 'Zophie', 'age': 7}
    print('name' in spam.keys())
    print('Zophie' in spam.values())
    print('color' in spam.keys())
    print('color' not in spam.keys())
    print('color' in spam)
    print('age' in spam)                # True
    print('Zophie' in spam)             # False
    print('Zophie' in spam.values())

def f69():
    # Dictionary get() method

    picnicItems = {'apples': 5, 'cups': 2}
    print('I am bringing ' + str(picnicItems.get('cups', 0)) + ' cups.')
    print('I am bringing ' + str(picnicItems.get('eggs', 0)) + ' eggs.')
    print('I am bringing ' + str(picnicItems.get('eggs', 'no value')) + ' eggs.')

    '''
    get() returns the value of the key, if the key exists.
    Because there is no 'eggs' key in the picnicItems dictionary,
    the default value 0 is returned by the get() method. Without
    using get() , the code would have caused an KeyError error.
    '''

def f70():
    # Dictionary setdefault() Method
    # sets a default value for key if key does not already have a value

    spam = {'name': 'Pooka', 'age': 5}
    if 'color' not in spam:
        spam['color'] = 'black'

    # The nice shortcut is using setdefault() method:
    spam = {'name': 'Pooka', 'age': 5}
    spam.setdefault('color', 'black')     # returns black
    print(spam.setdefault('color', 'black'))
    print(spam)

    spam.setdefault('color', 'white')     # returns black
    print(spam.setdefault('color', 'white'))
    print(spam)

def f71():
    # Using setdefault() to count number of occurance of a
    # character in string.

    message = 'It was a bright cold day in April, and the clocks were striking thirteen.'
    count = {}

    for character in message:
        count.setdefault(character, 0)
        count[character] = count[character] + 1

    print(count)

    '''
    The setdefault() method call ensures that the key is in the
    count dictionary (with a default value of 0 )

    So the program doesn’t throw a KeyError error,
    when count[character] = count[character] + 1 is executed.

    '''

def f72():
    # Pretty Printing - using pprint module

    # Following code counts number of the occurance of
    # character in the string and then using pprint to
    # display cleaner the items in a dictionary

    import pprint
    message = 'It was a bright cold day in April, and the clocks were striking thirteen.'
    count = {}

    for character in message:
        count.setdefault(character, 0)
        count[character] = count[character] + 1

    pprint.pprint(count)
    # print even prettier for nested loops:
    # pprint.pprint(someDictionaryValue)
    # print(pprint.pformat(someDictionaryValue))

def f73():
    # TIC-TOC-TOE - using dictionary to model the board
    theBoard = {'top-L': ' ', 'top-M': ' ', 'top-R': ' ', 'mid-L': ' ', 'mid-M': ' ',
                'mid-R': ' ', 'low-L': ' ', 'low-M': ' ', 'low-R': ' '}

    def printBoard(board):
        print(board['top-L'] + '|'+ board['top-M'] + '|' + board['top-R'])
        print('-+-+-')
        print(board['mid-L'] + '|'+ board['mid-M'] + '|' + board['mid-R'])
        print('-+-+-')
        print(board['low-L'] + '|'+ board['low-M'] + '|' + board['low-R'])
    turn = 'X'
    for i in range(9):
        printBoard(theBoard)
        print('Turn for ' + turn + '. Move on which space?')
        move = input()
        if move == 'q':
            return
        theBoard[move] = turn
        if turn == 'X':
            turn = 'O'
        else:
            turn = 'X'

    printBoard(theBoard)

def f74():
    # Nested Dictionaries and Lists
    # more complex data structure

    allGuests = {'Alice': {'apples': 5, 'pretzels': 12},
                 'Bob': {'ham sandwiches': 3, 'apples': 2},
                'Carol': {'cups': 3, 'apple pies': 1}}

    def totalBrought(guests, item):
        numBrought = 0
        for k, v in guests.items():
            numBrought = numBrought + v.get(item, 0)
        return numBrought

    print('Number of things being brought:')
    print(' - apples         ' + str(totalBrought(allGuests, 'apples')))
    print(' - Cups           ' + str(totalBrought(allGuests, 'cups')))
    print(' - Cakes          ' + str(totalBrought(allGuests, 'cakes')))
    print(' - Ham Sandwiches ' + str(totalBrought(allGuests, 'ham sandwiches')))
    print(' - Apple Pies     ' + str(totalBrought(allGuests, 'apple pies')))

    '''Inside the totalBrought() function, the for loop iterates over the key-
    value pairs in guests. Inside the loop, the string of the guest’s name is
    assigned to k , and the dictionary of picnic items they’re bringing is assigned
    to v . If the item parameter exists as a key in this dictionary, it’s value (the
    quantity) is added to numBrought. If it does not exist as a key, the get()
    method returns 0 to be added to numBrought .'''

def f75():
    # Fantasy Game Inventory
    print('project to do')

def f76():
    # List to Dictionary Function for Fantasy Game Inventory
    print('project to do')

def f77():
    # Chapter 6 - Manipulating Strings
    print('''Chapter 6 - Manipulating Strings:
    . String Literals
    . Indexing and Slicing Strings
    . The in and not in Operators with Strings
    . The upper(), lower(), isupper(), and islower() String Methods
    . The isX String Methods
    . The startswith() and endswith() String Methods
    . The join() and split() String Methods
    . Justifying Text with rjust(), ljust(), and center()
    . Removing Whitespace with strip(), rstrip(), and lstrip()
    . Copying and Pasting Strings with the pyperclip Module
    ''')

def f78():
    # Working with Strings - Double Quotes

    # By using double quotes, strings can have a single quote,
    # inside of string, and no need to be escaped.(interpolates)
    spam = "That is Alice's cat."
    print(spam)

    # When using single quotes, the single quote character,
    # inside string has to be escaped.
    spam = 'That is Alice\'s cat.'
    print(spam)

def f79():
    # Raw Strings and escape characters

    print('Printing string with special characters:')
    print("Hello there!\nHow are you?\nI\'m doing fine.")

    # You can place an r before the beginning quotation mark of a string to make
    # it a raw string. A raw string completely ignores all escape characters and
    # prints any backslash that appears in the string.
    print(r'That is Carol\'s cat.')

    es = r'''Escape ccharacters
        \' Single quote
        \" Double quote
        \t Tab
        \n Newline (line break)
        \\ Backslash'''

    print(es)

def f80():
    # Multiline Strings with Triple Quotes

    # Any quotes, tabs, or newlines in between the “triple quotes”,
    # are considered part of the string.
    print('''Dear Alice,

    Eve's cat has been arrested for catnapping, cat burglary, and extortion.

    Sincerely,
    Bob''')
    print()
    print("""Dear Alice,

    Eve's cat has been arrested for catnapping, cat burglary, and extortion.

    Sincerely,
    Bob""")
    print('Dear Alice,\n\nEve\'s cat has been arrested for catnapping, cat burglary, and extortion.\n\nSincerely,\nBob')

def f81():
    # Multiline Comments

    """This is a test Python program.
    Written by tester ...
    This program was designed for Python 3, not Python 2.
    """
    def spam():
        """This is a multiline comment to help
        explain what the spam() function does."""
        print('Hello!')

    spam()

def f82():
    # Indexing and Slicing Strings
    spam = 'Hello world!'
    print(spam[0])
    print(spam[4])
    print(spam[-1])
    print(spam[0:5])
    print(spam[:5])
    print(spam[6:])
    print(spam[0:5])

def f83():
    # The in and not in Operators with Strings

    spam = 'Hello world!'
    print('Hello' in 'Hello World')
    print('Hello' in 'Hello')
    print('HELLO' in 'Hello World')
    print('' in 'spam')
    print('cats' not in 'cats and dogs')

def f84():
    # String Methods: upper(), lower(), isupper(), and islower()

    spam = 'Hello world!'
    spam = spam.upper()
    print(spam)

    spam = spam.lower()
    print(spam)

    print('How are you?')
    feeling = input()
    if feeling.lower() == 'great':
        print('I feel great too.')
    else:
        print('I hope the rest of your day is good.')

    ''' Sample output
    How are you?
    GREat
    I feel great too.
    '''

def f85():
    # String Methods: isupper(), and islower()

    spam = 'Hello world!'
    print(spam.islower())
    print(spam.isupper())
    print('HELLO'.isupper())
    print('abc12345'.islower())
    print('12345'.islower())
    print('12345'.isupper())

    # upper and lower returns strings
    print('Hello'.upper())
    print('Hello'.upper().lower())
    print('Hello'.upper().lower().upper())
    print('HELLO'.lower())
    print('HELLO'.lower().islower())

def f86():
    # The isX String Methods

    print(
        # returns True if the string consists only of letters
    # and is not blank.
    'hello'.isalpha(),
    'hello123'.isalpha(),

    # isalnum() returns True if the string consists only of
    # letters and numbers
    'hello123'.isalnum(),
    'hello'.isalnum(),

    # isdecimal() returns True if the string consists only of numeric
    # characters and is not blank.
    '123'.isdecimal(),

    # isspace() returns True if the string consists only of spaces,
    # tabs, and new-lines and is not blank.
    '    '.isspace(),

    # istitle() returns True if the string consists only of words that
    # begin with an uppercase letter followed by only lowercase letters.
    'This Is Title Case'.istitle(),
    'This Is Title Case 123'.istitle(),
    'This Is not Title Case'.istitle(),
    'This Is NOT Title Case Either'.istitle(),
    )

def f87():
    # isX methods helps to validate user input
    while True:
        print('Enter your age:')
        age = input()
        if age.isdecimal():
            break
        print('Please enter a number for your age.')

    while True:
        print('Select a new password (letters and numbers only):')
        password = input()
        if password.isalnum():
            break
        print('Passwords can only have letters and numbers.')


''' Sample output
Enter your age:
forty two
Please enter a number for your age.
Enter your age:
42
Select a new password (letters and numbers only):
secr3t!
Passwords can only have letters and numbers.
Select a new password (letters and numbers only):
secr3t
'''

def f88():
    # startswith() and endswith() String Methods
    print(
        'Hello world!'.startswith('Hello'),
    # True
    'Hello world!'.endswith('world!'),
    # True
    'abc123'.startswith('abcdef'),
    # False
    'abc123'.endswith('12'),
    # False
    'Hello world!'.startswith('Hello world!'),
    # True
    'Hello world!'.endswith('Hello world!'),
    # True
    )

def f89():
    # join() and split() String Methods

    print(', '.join(['cats', 'rats', 'bats']))
    print(' '.join(['My', 'name', 'is', 'Simon']))
    print('ABC'.join(['My', 'name', 'is', 'Simon']))

    print('My name is Simon'.split())
    print('MyABCnameABCisABCSimon'.split('ABC'))
    print('My name is Simon'.split('m'))

def f90():
    # split a multiline string

    spam = '''Dear Alice,
    How have you been? I am fine.
    There is a container in the fridge
    that is labeled "Milk Experiment".
    Please do not drink it.
    Sincerely,
    Bob'''

    print(spam.split('\n'))

def f91():
    # Justifying Text with rjust(), ljust(), and center()

    print('Hello'.rjust(10))
    print('Hello'.rjust(20))
    print('Hello World'.rjust(20))
    print('Hello'.ljust(10))
    print('Hello'.rjust(20, '*'))
    print('Hello'.ljust(20, '-'))
    print('Hello'.center(20))
    print('Hello'.center(20, '='))

def f92():
    # Using Justifying Text for Tabular Data

    def printPicnic(itemsDict, leftWidth, rightWidth):
        print('PICNIC ITEMS'.center(leftWidth + rightWidth, '-'))
        for k, v in itemsDict.items():
            print(k.ljust(leftWidth, '.') + str(v).rjust(rightWidth))

    picnicItems = {'sandwiches': 4, 'apples': 12, 'cups': 4, 'cookies': 8000}
    printPicnic(picnicItems, 12, 5)
    printPicnic(picnicItems, 20, 6)

def f93():
    # Removing Whitespace with strip(), rstrip(), and lstrip()

    spam = '    Hello World     '
    print(spam.strip())
    print(spam.lstrip())
    print(spam.rstrip())

    spam = 'SpamSpamBaconSpamEggsSpamSpam'
    print(spam.strip('ampS'))
    print(spam.strip('Spma'))
    # Above two are  the same, order of argument character is ignored.

def f94():
    # Copy and Paste Strings with the pyperclip Module

    import pyperclip
    pyperclip.copy('Hello world!')
    print(pyperclip.paste())

def f95():
    # An insecure password locker program.

    import pyperclip

    # Step 1: Program Design and Data Structures
    PASSWORDS = {'email': 'F7minlBDDuvMJuxESSKHFhTxFtjVB6',
                 'blog': 'VmALvQyKAxiVH5G8v01if1MLZF3sdt',
                'luggage': '12345'}


    # Step 2: Handle Command Line Arguments
    import sys
    if len(sys.argv) < 2:
        print('Usage: python pw.py [account] - copy account password')
        sys.exit()

    account = sys.argv[1]  # first command line arg is the account name


    # Step 3: Copy the Right Password
    if account in PASSWORDS:
        pyperclip.copy(PASSWORDS[account])
        print('Password for ' + account + ' copied to clipboard.')
    else:
        print('There is no account named ' + account)

def f96():
    # Adding Bullets to Wiki Markup

    # Step 1: Copy and Paste from the Clipboard
    import pyperclip

    text = pyperclip.paste()

    # Step 2 - # TODO: Separate lines and add stars.
    lines = text.split('\n')
    for i in range(len(lines)):    # loop through all indexes in the "lines" list
        lines[i] = '* ' + lines[i] # add star to each string in "lines" list

    # Step 3: Join the Modified Lines
    text = '\n'.join(lines)
    pyperclip.copy(text)

    # When this program is run, it replaces the text on the clipboard with
    # text that has stars at the start of each line.

def f97():
    # Chapter 7 - Pattern Matching With Reular Expression
    print('')
def f98():
    # Finding Patterns of Text Without Regular Expressions

    def isPhoneNumber(text):
        if len(text) != 12:
            return False
        for i in range(0, 3):   # i = [0, 1, 2]
            if not text[i].isdecimal():   # Convert string to list and
                                            # parsing the list for decimals
                return False
        if text[3] != '-':
            return False
        for i in range(4, 7):    # i = [4, 5, 6]
            if not text[i].isdecimal():
                return False
        if text[7] != '-':
            return False
        for i in range(8, 12):   # i = [8, 9, 10, 11]
            if not text[i].isdecimal():
                return False
        return True

    print('415-555-4242 is a phone number:')
    print(isPhoneNumber('415-555-4242'))
    print('Moshi moshi is a phone number:')
    print(isPhoneNumber('Moshi moshi'))

def f99():
    # Finding pattern without RE from a string

    def isPhoneNumber(text):
        if len(text) != 12:
            return False
        for i in range(0, 3):
            if not text[i].isdecimal():
                return False
        if text[3] != '-':
            return False
        for i in range(4, 7):
            if not text[i].isdecimal():
                return False
        if text[7] != '-':
            return False
        for i in range(8, 12):
            if not text[i].isdecimal():
                return False
        return True

    message = 'Call me at 415-555-1011 tomorrow. 415-555-9999 is my office.'
    for i in range(len(message)):
        chunk = message[i:i+12]   # parsing chunks of 12 character in string
        if isPhoneNumber(chunk):
            print('Phone number found: ' + chunk)

    print('Done')

def f100():
    # Matching regex objects

    # In a raw string a backslash is taken as meaning just as backslash not as
    # a escape sequence. if we donot use r'', then we have to escape all the
    # backslashes between the single quotes

    import re

    phoneNumRegex = re.compile(r'\d\d\d-\d\d\d-\d\d\d\d')
    mo = phoneNumRegex.search('My number is 415-555-4242.')
    print('Phone number found: ' + mo.group())

'''
Steps in using regex:
---------------------
1.   Import the regex module with import re.
2.   Create a Regex object with the re.compile() function.
     (Remember to use a raw string.)
3.   Pass the string you want to search into the Regex object’s search() method.
     This returns a Match object.
4.   Call the Match object’s group() method to return a string of the actual
     matched text.
'''

def f101():
    # Pattern matching with grouping

    import re

    phoneNumRegex = re.compile(r'(\d\d\d)-(\d\d\d-\d\d\d\d)')
    mo = phoneNumRegex.search('My number is 415-555-4242.')
    print(mo.group(1))
    print(mo.group(2))
    print(mo.group(0))  # returns entire match
    print(mo.group())   # returns entire match, the same as group(0)
    print(mo.groups())  # Retrives all groups, attention at plural form of group
    areaCode, mainNumber = mo.groups()  # group 1 and 2 are elements in a list
                                        # generated by method groups()
    print(areaCode)
    print(mainNumber)

def f102():
    # Pattern matching with grouping and escaping parantheses

    import re

    phoneNumRegex = re.compile(r'(\(\d\d\d\)) (\d\d\d-\d\d\d\d)')
    mo = phoneNumRegex.search('My phone number is (415) 555-4242.')
    print(mo.group(1))
    print(mo.group(2))

def f103():
    # Matching multiple groups with the pipe

    import re

    heroRegex = re.compile (r'Batman|Tina Fey')
    mo1 = heroRegex.search('Batman and Tina Fey.')
    print(mo1.group())

    mo2 = heroRegex.search('Tina Fey and Batman.')
    print(mo2.group())

def f104():
    # Matching groups with the pipe as part of re

    import re

    batRegex = re.compile(r'Bat(man|mobile|copter|bat)')
    mo = batRegex.search('Batmobile lost a wheel')
    print(mo.group())      # returns all full-matched text 'Batmobile'
    print(mo.group(1))     # returns the first match it founds,
                            # in first group

def f105():
    # Optional matching with the question mark

    import re

    batRegex = re.compile(r'Bat(wo)?man')
    mo1 = batRegex.search('The Adventures of Batman')
    print(mo1.group())

    mo2 = batRegex.search('The Adventures of Batwoman')
    print(mo2.group())

    # (wo)? means that the pattern wo is an optional group.
    # the re matches zero instances or one instance of wo in it.

    # same way for phone numbers that do or do not have an area code:
    phoneRegex = re.compile(r'(\d\d\d-)?\d\d\d-\d\d\d\d')
    mo1 = phoneRegex.search('My number is 415-555-4242')
    print(mo1.group())

    mo2 = phoneRegex.search('My number is 555-4242')
    print(mo2.group())

    # (pattern)?
    # Match zero or one of the group pre-ceding this question mark.

def f106():
    # Matching zero or more with star

    import re

    batRegex = re.compile(r'Bat(wo)*man')
    mo1 = batRegex.search('The Adventures of Batman')
    print(mo1.group())

    mo2 = batRegex.search('The Adventures of Batwoman')
    print(mo2.group())

    mo3 = batRegex.search('The Adventures of Batwowowowoman')
    print(mo3.group())

    # (pattern)* means match zero or more of the group
    # that precedes the star

def f107():
    # Matching one or more with the plus
    # the group preceding a plus must appear at least once.

    import re

    batRegex = re.compile(r'Bat(wo)+man')
    mo1 = batRegex.search('The Adventures of Batwoman')
    print(mo1.group())


    mo2 = batRegex.search('The Adventures of Batwowowowoman')
    print(mo2.group())

    mo3 = batRegex.search('The Adventures of Batman')
    print(mo3 == None)

def f108():
    # Matching specific repetitions with {}
    # (pattern){a_number}
    import re

    haRegex = re.compile(r'(Ha){3}')
    mo1 = haRegex.search('HaHaHa')
    print(mo1.group())

    mo2 = haRegex.search('Ha')
    print(mo2 == None)

def f109():
    # Greedy and nongreedy matching
    # RE by default is greedy, matches the longest match.
    # To have non-greedy match use {}?

    import re

    greedyHaRegex = re.compile(r'(Ha){3,5}')  # greedy
    mo1 = greedyHaRegex.search('HaHaHaHaHa')
    print(mo1.group())

    nongreedyHaRegex = re.compile(r'(Ha){3,5}?')  # non-greedy
    mo2 = nongreedyHaRegex.search('HaHaHaHaHa')
    print(mo2.group())

    # question mark can have two different meanings in RE:
    # declaring a nongreedy match or flagging an optional group.

def f110():
    # RE findall() Method
    # While search() will return a Match object of the first matched text
    # in the searched string, the findall() method will return the strings
    # of every match in the searched string.

    import re

    phoneNumRegex = re.compile(r'\d\d\d-\d\d\d-\d\d\d\d')
    mo = phoneNumRegex.search('Cell: 415-555-9999 Work: 212-555-0000')
    print(mo.group())

    phoneNumRegex = re.compile(r'\d\d\d-\d\d\d-\d\d\d\d') # has no groups
    print(phoneNumRegex.findall('Cell: 415-555-9999 Work: 212-555-0000'))

    # findall() will not return a Match object but a list of strings
    # as long as there are no groups in the regular expression.

    # When groups are in RE, then findall returns list of tuples
    phoneNumRegex = re.compile(r'(\d\d\d)-(\d\d\d)-(\d\d\d\d)') # has groups
    print(phoneNumRegex.findall('Cell: 415-555-9999 Work: 212-555-0000'))

def f111():
    # RE Character Classes
    # Character classes are nice for shortening REs
    # \d  is equal to (0|1|2|3|4|5|6|7|8|9)
    # \D Any character that is not a numeric digit from 0 to 9.
    # \w Any letter, numeric digit, or the underscore character. (word character)
    # \W Any character that is not a letter, numeric digit, or the underscore character.
    # \s Any space, tab, or newline character.(space character)
    # \S Any character that is not a space, tab, or newline.

    import re

    xmasRegex = re.compile(r'\d+\s\w+')
    f = xmasRegex.findall('12 drummers, 11 pipers, 10 lords, 9 ladies, 8 maids, swans, 6 geese, 5 rings, 4 birds, 3 hens, 2 doves, 1 partridge')
    print(f)

def f112():
    # Making your own character classes
    # match any vowel, both lowercase and uppercase

    import re

    vowelRegex = re.compile(r'[aeiouAEIOU]')
    print(vowelRegex.findall('RoboCop eats baby food. BABY FOOD.'))

    # [a-zA-Z0-9] will match all lowercase letters,
    # uppercase letters, and numbers

    # negative character class by ^
    consonantRegex = re.compile(r'[^aeiouAEIOU]')
    print(consonantRegex.findall('RoboCop eats baby food. BABY FOOD.'))

def f113():
    # Caret and dollar sign characters

    import re

    beginsWithHello = re.compile(r'^Hello')
    print(beginsWithHello.search('Hello world!'))
    print(beginsWithHello.search('He said hello.') == None)

    endsWithNumber = re.compile(r'\d$')
    print(endsWithNumber.search('Your number is 42'))
    print(endsWithNumber.search('Your number is forty two.') == None)

    # To match strings that both begin and end with one or
    # more numeric characters:
    wholeStringIsNum = re.compile(r'^\d+$')
    print(wholeStringIsNum.search('1234567890'))
    print(wholeStringIsNum.search('12345xyz67890') == None)
    print(wholeStringIsNum.search('12 34567890') == None)

def f114():
    # Wild card character
    # The . (or dot) character in a RE is called a wildcard
    # and will match just one character except for a newline.

    import re

    atRegex = re.compile(r'.at')
    print(atRegex.findall('The cat in the hat sat on the flat mat.'))

def f115():
    # Matching everything with dot-star

    import re

    nameRegex = re.compile(r'First Name: (.*) Last Name: (.*)')
    mo = nameRegex.search('First Name: Al Last Name: Sweigart')
    print(mo.group(1))
    print(mo.group(2))

    # The dot-star uses greedy mode,
    # to match any and all text in a nongreedy fashion, use (.*?)

    nongreedyRegex = re.compile(r'<.*?>')
    mo = nongreedyRegex.search('<To serve man> for dinner.>')
    print(mo.group())

    greedyRegex = re.compile(r'<.*>')
    mo = greedyRegex.search('<To serve man> for dinner.>')
    print(mo.group())

def f116():
    # Matching newlines with the dot character

    import re

    # To match everything only up to the first newline character,
    # newline character is not encluded
    noNewlineRegex = re.compile('.*')
    print(noNewlineRegex.search(
        'Serve the public trust.\nProtect the innocent.\nUphold the law.'
    ).group())

    # re.DOTALL matches everything,
    # To include newline character:
    newlineRegex = re.compile('.*', re.DOTALL)
    print(newlineRegex.search(
        'Serve the public trust.\nProtect the innocent.\nUphold the law.'
    ).group())

def f117():
    # Case-insensitive matching
    # Use re.I or re.IGNORECASE

    import re

    robocop = re.compile(r'robocop', re.I)
    print(robocop.search(
        'RoboCop is part man, part machine, all cop.').group())
    print(robocop.search(
        'ROBOCOP protects the innocent.').group())

    print(robocop.search(
        'Al, why does your programming book talk about robocop so much?'
    ).group())

def f118():
    # Substituting strings with sub() method

    import re

    namesRegex = re.compile(r'Agent \w+')
    print(namesRegex.sub(
        'CENSORED', 'Agent Alice gave the secret documents to Agent Bob.'))

    # using with groups: text of group 1 , 2 , 3
    agentNamesRegex = re.compile(r'Agent (\w)\w*')
    print(agentNamesRegex.sub(r'\1****', 'Agent Alice told Agent Carol \
    that Agent Eve knew Agent Bob was a double agent.'))

def f119():
    # Managing complex regexes
    # You can simplify the hard-to-read REs by using verbose mode
    # re.VERBOSE ignores whitespace and comments inside the RE string.

    import re

    # Hard to read :
    phoneRegex = re.compile(r'((\d{3}|\(\d{3}\))?(\s|-|\.)?\d{3}(\s|-|\.)\d{4}(\s*(ext|x|ext.)\s*\d{2,5})?)')

    # Expand to multiline using comments:
    phoneRegex = re.compile(r'''(
                            (\d{3}|\(\d{3}\))?              # area code
                            (\s|-|\.)?                      # separator
                            \d{3}                           # first 3 digits
                            (\s|-|\.)                       # separator
                            \d{4}                           # last 4 digits
                            (\s*(ext|x|ext.)\s*\d{2,5})?    # extension
                            )''', re.VERBOSE)


    someRegexValue = re.compile('foo', re.IGNORECASE | re.DOTALL)
    someRegexValue = re.compile('foo', re.IGNORECASE | re.DOTALL | re.VERBOSE)

def f120():
    # Combining re.IGNORECASE, re.DOTALL, and re.VERBOSE

    import re

    someRegexValue = re.compile('foo', re.IGNORECASE | re.DOTALL)
    someRegexValue = re.compile('foo', re.IGNORECASE | re.DOTALL | re.VERBOSE)

def f121():
    # Finds phone numbers and email addresses on the clipboard

    import re, pyperclip

    # Step 1: Create a Regex for Phone Numbers
    phoneRegex = re.compile(r'''(
                            (\d{3}|\(\d{3}\))?              # area code
                            (\s|-|\.)?                      # separator
                            (\d{3})                         # first 3 digits
                            (\s|-|\.)                       # separator
                            (\d{4})                         # last 4 digits
                            (\s*(ext|x|ext.)\s*(\d{2,5}))?  # extension
                            )''', re.VERBOSE)

    # Step 2: Create a Regex for Email Addresses
    emailRegex = re.compile(r'''(
                            [a-zA-Z0-9._%+-]+               # username
                            @                               # @ symbol
                            [a-zA-Z0-9.-]+                  # domain name
                            (\.[a-zA-Z]{2,4})               # dot-something
                            )''', re.VERBOSE)

    # Step 3: Find matches in clipboard text.
    text = str(pyperclip.paste())
    matches = []
    for groups in phoneRegex.findall(text):
        phoneNum = '-'.join([groups[1], groups[3], groups[5]])
        if groups[8] != '':
            phoneNum += ' x' + groups[8]
        matches.append(phoneNum)
    for groups in emailRegex.findall(text):
        matches.append(groups[0])

    # Step 4: Copy results to the clipboard.
    if len(matches) > 0:
        pyperclip.copy('\n'.join(matches))
        print('Copied to clipboard:')
        print('\n'.join(matches))
    else:
        print('No phone numbers or email addresses found.')

def f122():
    # In this assignment you will read through  and  parse  a file
    # with text and numbers. You will extract all the numbers in -
    # the file and compute the sum of the numbers.

    import re

    fileName = input('Enter file name: ')
    fh = open(fileName)

    # Initiate an empty list and extend it with list of
    # numbers found in the file.
    allNumbers = []

    for line in fh:
        line = line.rstrip()

        # Restart the loop if line has no numbers:
        if re.search('[0-9]+', line) == None : continue

        # extract numbers from each line, ignoring zero/s for
        # numbers starting with zero:
        numbers = re.findall('0*([0-9]+)', line)

        # extend the list allNumbers with list of numbers,
        # extracted from each line
        allNumbers.extend(numbers)

    # Compute the sum of all elements of the list,
    # that includes all numbers in the file:
    total = 0
    for n in allNumbers:
        total = total + int(n)

    print(total)

def f123():
    # Chapter 8 - Reading and Writing Files
    print('Chapter - Reading and Writing Files')

def f124():
    # Create strings for file names

    import os
    # os.path is a module inside the os module

    print(os.path.join('usr', 'bin', 'spam'))
    # On Windows the ouput will be with backslash

    myFiles = ['accounts.txt', 'details.csv', 'invite.docx']
    for filename in myFiles:
        print(os.path.join('/home/username/asweigart', filename))

def f125():
    # The Current Working Directory

    import os

    print(os.getcwd())
    os.chdir('/var/log')         # returns None
    print(os.getcwd())

    # After  running this function you get error in main program
    # because the path to current running script has been changed.

def f126():
    # Creating New Folders with os.makedirs()

    import os
    os.makedirs('/tmp/delicious/walnut/waffles')
    os.system('ls /tmp/delicious/walnut/')

    # You get error if you run this function forthe second time
    # because the directory already exists.

def f127():
    # Handling Absolute and Relative Paths

    import os
    print(os.path.abspath('.'))
    print(os.path.abspath('./Scripts'))
    print(os.path.isabs('.'))
    print(os.path.isabs(os.path.abspath('.')))

    print(os.path.relpath('/var/log', '/var/ftp'))
    print(os.getcwd())

    path = '/windows/logs/messages'
    print(os.path.basename(path))
    print(os.path.dirname(path))

    calcFilePath = '/Windows/System32/calc.exe'
    print(os.path.split(calcFilePath))
    print((os.path.dirname(calcFilePath), os.path.basename(calcFilePath)))

    print('/usr/bin'.split(os.path.sep))

def f128():
    # Finding File Sizes and Folder Contents

    import os
    print(os.path.getsize('runme.py'))
    print(os.listdir('/home'))

    # Calculating total size in current folder
    totalSize = 0
    for filename in os.listdir('./'):
        totalSize = totalSize + os.path.getsize(os.path.join('./', filename))

    print(totalSize)

def f129():
    # Checking Path Validity

    import os
    print(os.path.exists('/var/log'))
    print(os.path.exists('/tmp/some_made_up_folder'))
    print(os.path.isdir('/var/log'))
    print(os.path.isfile('/var/log'))
    print(os.path.isdir('/var/log/messages'))
    print(os.path.isfile('/var/log/messages'))

    # Check if dvd drive exists
    print(os.path.exists('D:\\'))

def f130():
    # The File Reading/Writing Process

    # open() function returns a File object, and opens file in plain text mode.
    # Assuming file hello.txt already created.

    helloFile = open('./hello.txt')
    # You also can open a file in read mode with 'r' option,
    # the following is the same as above:
    # helloFile = open('./hello.txt', 'r')

    helloContent = helloFile.read()
    print(helloContent)

    # Use the readlines() method to get a list of string,
    # values from the file, one string for each line of text.

    relFile = open('/etc/os-release')
    print(relFile.readlines())

def f131():
    # Writing to Files

    baconFile = open('bacon.txt', 'w')    # open in writr mode
    baconFile.write('Hello world!\n')
    baconFile.close()

    baconFile = open('bacon.txt', 'a')
    baconFile.write('Bacon is not a vegetable.')
    baconFile.close()

    baconFile = open('bacon.txt')
    content = baconFile.read()
    baconFile.close()

    print(content)

def f132():
    # Saving Variables with the shelve Module

    import shelve

    # shelf value is stored in shelfFile:
    shelfFile = shelve.open('mydata')
    cats = ['Zophie', 'Pooka', 'Simon']

    # You can make changes to the shelf value as if it were a dictionary.
    shelfFile['cats'] = cats
    shelfFile.close()


    shelfFile = shelve.open('mydata')
    print(type(shelfFile))
    print(shelfFile['cats'])
    shelfFile.close()

    shelfFile = shelve.open('mydata')
    print(list(shelfFile.keys()))
    print(list(shelfFile.values()))
    shelfFile.close()

    # but if you want to save data from your Python programs,
    # use the shelve module

def f133():
    # Saving Variables with the pprint.pformat() Function

    '''
    Say you have a dictionary stored in a variable and you want to save this
    variable and its contents for future use. Using pprint.pformat() will give
    you a string that you can write to .py file. This file will be your very
    own module that you can import when-ever you want to use the variable stored in it.
    '''

    import pprint

    cats = [{'name': 'Zophie', 'desc': 'chubby'}, {'name': 'Pooka', 'desc': 'fluffy'}]
    pprint.pformat(cats)

    fileObj = open('myCats.py', 'w')
    fileObj.write('cats = ' + pprint.pformat(cats) + '\n')
    fileObj.close()

    ''' In terminal:
    $ cat myCats.py
    cats = [{'desc': 'chubby', 'name': 'Zophie'}, {'desc': 'fluffy', 'name': 'Pooka'}]
    '''

    '''
    And since Python scripts are themselves just text files with the .py file
    extension, your Python programs can even generate other Python programs.
    You can then import these files into scripts.'''

    import myCats
    print(myCats.cats)
    print(myCats.cats[0])
    print(myCats.cats[0]['name'])

def f134():
    # Generating Random Quiz Files

    import random

    # Step 1: Store the Quiz Data in a Dictionary
    capitals = {
        'Alabama': 'Montgomery', 'Alaska': 'Juneau', 'Arizona': 'Phoenix',
    'Arkansas': 'Little Rock', 'California': 'Sacramento', 'Colorado': 'Denver',
    'Connecticut': 'Hartford', 'Delaware': 'Dover', 'Florida': 'Tallahassee',
    'Georgia': 'Atlanta', 'Hawaii': 'Honolulu', 'Idaho': 'Boise', 'Illinois': 'Springfield',
    'Indiana': 'Indianapolis', 'Iowa': 'Des Moines', 'Kansas': 'Topeka',
    'Kentucky': 'Frankfort', 'Louisiana': 'Baton Rouge', 'Maine': 'Augusta',
    'Maryland': 'Annapolis', 'Massachusetts': 'Boston', 'Michigan': 'Lansing',
    'Minnesota': 'Saint Paul', 'Mississippi': 'Jackson', 'Missouri': 'Jefferson City',
    'Montana': 'Helena', 'Nebraska': 'Lincoln', 'Nevada': 'Carson City',
    'New Hampshire': 'Concord', 'New Jersey': 'Trenton', 'New Mexico': 'Santa Fe',
    'New York': 'Albany', 'North Carolina': 'Raleigh', 'North Dakota': 'Bismarck',
    'Ohio': 'Columbus', 'Oklahoma': 'Oklahoma City', 'Oregon': 'Salem',
    'Pennsylvania': 'Harrisburg', 'Rhode Island': 'Providence',
    'South Carolina': 'Columbia', 'South Dakota': 'Pierre', 'Tennessee': 'Nashville',
    'Texas': 'Austin', 'Utah': 'Salt Lake City', 'Vermont': 'Montpelier',
    'Virginia': 'Richmond', 'Washington': 'Olympia', 'West Virginia': 'Charleston',
    'Wisconsin': 'Madison', 'Wyoming': 'Cheyenne'}

    # Step 2: Create the Quiz File and Shuffle the Question Order
    for quizNum in range(35):
        # Create the quiz and answer key files.
        quizFile = open('capitalsquiz%s.txt' % (quizNum + 1), 'w')
        answerKeyFile = open('capitalsquiz_answers%s.txt' % (quizNum + 1), 'w')

        # Write out the header for the quiz.
        quizFile.write('Name:\n\nDate:\n\nPeriod:\n\n')
        quizFile.write((' ' * 20) + 'State Capitals Quiz (Form %s)' % (quizNum + 1))
        quizFile.write('\n\n')

        # Shuffle the order of the states.
        states = list(capitals.keys())
        random.shuffle(states)

    # Step 3: Create the Answer Options
    # Loop through all 50 states, making a question for each.
    for questionNum in range(50):
        # Get right and wrong answers.
        correctAnswer = capitals[states[questionNum]]
        wrongAnswers = list(capitals.values())
        del wrongAnswers[wrongAnswers.index(correctAnswer)]
        wrongAnswers = random.sample(wrongAnswers, 3)
        answerOptions = wrongAnswers + [correctAnswer]
        random.shuffle(answerOptions)


    # Step 4: Write Content to the Quiz and Answer Key Files
        quizFile.write('%s. What is the capital of %s?\n' % (questionNum + 1, states[questionNum]))
        for i in range(4):
            quizFile.write('    %s. %s\n' % ('ABCD'[i], answerOptions[i]))
        quizFile.write('\n')

        # Write the answer key to a file.
        answerKeyFile.write('%s. %s\n' % (questionNum + 1, 'ABCD'[answerOptions.index(correctAnswer)]))

    quizFile.close()
    answerKeyFile.close()

def f135():
    # Multiclipboard
    print('')
    # Step 1: Comments and Shelf Setup

def f136():
    # Chapter 9 - Organizing Files
    print('Chapter 9 - Organizing Files')

def f137():
    # Copying Files and Folders - using shutil module

    # shutil.copy(source, destination)
    # copy() returns a string of the path of the copied file.

    import shutil, os
    print(shutil.copy('./runme.py', '/tmp'))
    print(shutil.copy('./runme.py', '/tmp/runme2.py'))
    os.listdir('/tmp')

    # shutil.copytree(source, destination)
    print('--- Cope all files and subfolders from a folder by copytree() --- ')
    shutil.copytree('./', '/tmp/test')

    # Above copies all files and subfolders in current directory into /tmp/test,
    # the test directroy will be created if it is not already exist.

def f138():
    # Moving and Renaming Files and Folders

    ''' Examples:
    >>> shutil.move('C:\\bacon.txt', 'C:\\eggs\\new_bacon.txt')
    'C:\\eggs\\new_bacon.txt'
    >>> shutil.move('C:\\bacon.txt', 'C:\\eggs')
    'C:\\eggs'
    >>> shutil.move('spam.txt', 'c:\\does_not_exist\\eggs\\ham')

    If move() cannot find the destination folder or file,
    then will get FileNotFoundError error.
    '''
    print()

def f139():
    # Permanently Deleting Files and Folders

    '''
    .   Calling os.unlink(path) will delete the file at path.
    .   Calling os.rmdir(path) will delete the folder at path . This folder must be
        empty of any files or folders.
    .   Calling shutil.rmtree(path) will remove the folder at path , and all files
        and folders it contains will also be deleted.
    '''

    import os
    for filename in os.listdir('/tmp'):
        if filename.endswith('.py'):
            #os.unlink(filename)
            print(filename)

def f140():
    # Safe Deletes with the send2trash Module

    # install the 3rd party module send2trash, first:
    # sudo pip3 install send2trash

    import send2trash

    baconFile = open('bacon.txt', 'a')        # creates the file
    print(baconFile.write('Bacon is not a vegetable.'))
    baconFile.close()
    send2trash.send2trash('bacon.txt')

def f141():
    # Walk through the directory tree

    import os

    for folderName, subfolders, filenames in os.walk('/tmp'):
        print('The current folder is ' + folderName)
    for subfolder in subfolders:
        print('SUBFOLDER OF ' + folderName + ': ' + subfolder)

    for filename in filenames:
        print('FILE INSIDE ' + folderName + ': '+ filename)
    print('')

    # os.walk, doesn't return list of files/subfolders recursively, it returns
    # current folder name, subfolders and file names in current directory.

def f142():
    # Compressing Files with the zipfile Module

    import zipfile, os

    # First create a zip file, from alll files in current directory:
    os.system('zip -r example.zip ./')
    # 1st create a ZipFile object, similar to when creating a file object
    exampleZip = zipfile.ZipFile('example.zip')

    print('\nList of files and folders in example.zip :')
    print(exampleZip.namelist())

    fileInfo = exampleZip.getinfo('runme.py')
    print(fileInfo)
    print('--------')
    print(fileInfo.file_size)
    print(fileInfo.compress_size)
    print('Compressed file is %sx smaller!' % (round(fileInfo.file_size / fileInfo.compress_size, 2)))
    exampleZip.close()
    os.unlink('example.zip')

def f143():
    # Extracting from ZIP Files

    import zipfile, shutil, os

    os.makedirs('/tmp/test')
    os.system('zip -r /tmp/test/example.zip ./')   # Create a zip file in /tmp
    # os.chdir('/tmp/test') # move to the folder with example.zip
    exampleZip = zipfile.ZipFile('/tmp/test/example.zip')  # Creating ZipFile object
    exampleZip.extractall()
    exampleZip.close()
    print(os.listdir('/tmp/test'))
    shutil.rmtree('/tmp/test')

def f144():
    # Creating and Adding to ZIP Files

    import zipfile, os

    # Create an ZipFile object in write mode, similar to file object in write mode
    newZip = zipfile.ZipFile('new.zip', 'w')

    newZip.write('runme.py', compress_type=zipfile.ZIP_DEFLATED)
    newZip.close()

    os.system('ls')
    os.unlink('new.zip')

def f145():
    # Renaming Files with American-Style Dates to
    # European-Style Dates
    # shutil.move() function can be used to rename files

    # Step 1: Create a Regex for American-Style Dates
    import shutil, os, re

    # Create a regex that matches files with the American date format.
    datePattern = re.compile(r"""^(.*?) # all text before the date
                             ((0|1)?\d)-                     # one or two digits for the month
                             ((0|1|2|3)?\d)-                 # one or two digits for the day
                             ((19|20)\d\d)                   # four digits for the year
                             (.*?)$                          # all text after the date
                             """, re.VERBOSE)
    # Step 2: Identify the Date Parts from the Filenames
    # Loop over the files in the working directory.
    for amerFilename in os.listdir('.'):
        mo = datePattern.search(amerFilename)
        # Skip files without a date.
        if mo == None:
            continue
        # Get the different parts of the filename.
        beforePart = mo.group(1)
        monthPart  = mo.group(2)
        dayPart    = mo.group(4)
        yearPart   = mo.group(6)
        afterPart  = mo.group(8)

        # Step 3: Form the New Filename and Rename the Files
        # Form the European-style filename.
        euroFilename = beforePart + dayPart + '-' + monthPart + '-' + yearPart + afterPart

        # Get the full, absolute file paths.
        absWorkingDir = os.path.abspath('.')
        amerFilename = os.path.join(absWorkingDir, amerFilename)
        euroFilename = os.path.join(absWorkingDir, euroFilename)

        # Rename the files.
        print('Renaming "%s" to "%s"...' % (amerFilename, euroFilename))
        #shutil.move(amerFilename, euroFilename) # uncomment after testing

def f146():
    # Backing Up a Folder into a ZIP File

    import zipfile, os

    # Step 1: Figure Out the ZIP File’s Name
    def backupToZip(folder):
        # Backup the entire contents of "folder" into a ZIP file.
        folder = os.path.abspath(folder)    # make sure folder is absolute

        # Figure out the filename this code should use based on
        # what files already exist.
        number = 1
        while True:
            zipFilename = os.path.basename(folder) + '_' + str(number) + '.zip'
            if not os.path.exists(zipFilename):
                break
            number = number + 1

        # Step 2: Create the ZIP file.
        print('Creating %s...' % (zipFilename))
        backupZip = zipfile.ZipFile(zipFilename, 'w')


        # step 3: Walk the entire folder tree and compress the files
        # in each folder.
        for foldername, subfolders, filenames in os.walk(folder):
            print('Adding files in %s...' % (foldername))
            # Add the current folder to the ZIP file.
            backupZip.write(foldername)

            # Add all the files in this folder to the ZIP file.
            for filename in filenames:
                newBase = os.path.basename(folder) + '_'
                if filename.startswith(newBase) and filename.endswith('.zip'):
                    continue   # don't backup the backup ZIP files
                backupZip.write(os.path.join(foldername, filename))
            backupZip.close()
            print('Done.')


    backupToZip('/home/bmzi/bin')

def f147():
    # Chapter 10 - Debugging
    print('Chapter 10 - Debugging')

def f148():
    # Raising Exceptions

    def boxPrint(symbol, width, height):
        if len(symbol) != 1:
            raise Exception('Symbol must be a single character string.')
        if width <= 2:
            raise Exception('Width must be greater than 2.')
        if height <= 2:
            raise Exception('Height must be greater than 2.')

        print(symbol * width)
        for i in range(height - 2):
            print(symbol + (' ' * (width - 2)) + symbol)

        print(symbol * width)


    for sym, w, h in (('*', 4, 4), ('O', 20, 5), ('x', 1, 3), ('ZZ', 3, 3)):
        try:
            boxPrint(sym, w, h)
        except Exception as err:
            print('An exception happened: ' + str(err))


    # If there are no try and except statements covering the raise statement
    # that raised the exception, the program simply crashes and displays the
    # exception’s error message.

def f149():
    # Getting the Traceback as a String

    import traceback, os

    try:
        raise Exception('This is the error message.')
    except:
        errorFile = open('errorInfo.txt', 'w')
        errorFile.write(traceback.format_exc())
        errorFile.close()
        print('The traceback info was written to errorInfo.txt.')

    print('---')
    os.system('ls')
    print('---')
    os.system('cat errorInfo.txt')
    os.unlink('errorInfo.txt')

def f150():
    # Assertions - is a sanity check of the code
    # an assert statement says, 'I assert that this condition holds true,
    # and if not, there is a bug somewhere in the program.'

    ''' assert statement consists of the following:
    .   The assert keyword
    .   A condition (that is, an expression that evaluates to True or False )
    .   A comma
    .   A string to display when the condition is False'''

    # After running following code in a separate file like test.py:
    '''
    podBayDoorStatus = 'open'
    assert podBayDoorStatus == 'open', 'The pod bay doors need to be "open".'
    podBayDoorStatus = 'I\'m sorry, Dave. I\'m afraid I cant do that.\''
    assert podBayDoorStatus == 'open', 'The pod bay doors need to be "open".'
    '''

    # output error:
    '''
    Traceback (most recent call last):
      File "./test.py", line 7, in <module>
        assert podBayDoorStatus == 'open', 'The pod bay doors need to be "open".'
    AssertionError: The pod bay doors need to be "open".

    '''
    print()

def f151():
    # Using an Assertion in a Traffic Light Simulation

    def switchLights(stoplight):
        for key in stoplight.keys():
            if stoplight[key] == 'green':
                stoplight[key] = 'yellow'
            elif stoplight[key] == 'yellow':
                stoplight[key] = 'red'
            elif stoplight[key] == 'red':
                stoplight[key] = 'green'

    market_2nd = {'ns': 'green', 'ew': 'red'}
    mission_16th = {'ns': 'red', 'ew': 'green'}

    switchLights(market_2nd)

    # assert 'red' in stoplight.values(), 'Neither light is red! ' + str(stoplight)

    ''' If you uncomment the line above your program would crash with this error
    message:
    Traceback (most recent call last):
        File "carSim.py", line 14, in <module>
            switchLights(market_2nd)
        File "carSim.py", line 13, in switchLights
            assert 'red' in stoplight.values(), 'Neither light is red! ' + str(stoplight)
    AssertionError: Neither light is red! {'ns': 'yellow', 'ew': 'green'}
    '''

    # Assertions can be disabled by passing the -O option when running Python.

def f152():
    # Logging - Using the logging Module

    import logging
    logging.basicConfig(level=logging.DEBUG, format=' %(asctime)s - %(levelname)s - %(message)s')
    logging.debug('Start of program')

    def factorial(n):
        logging.debug('Start of factorial(%s%%)' % (n))
        total = 1
        for i in range(n + 1):
            total *= i
            logging.debug('i is ' + str(i) + ', total is ' + str(total))
        logging.debug('End of factorial(%s%%)' % (n))
        return total

    print(factorial(5))
    logging.debug('End of program')

    ''' About this code:
    Here, we use the logging.debug() function when we want to print log
    information. This debug() function will call basicConfig() , and a line of infor-
    mation will be printed. This information will be in the format we specified
    in basicConfig() and will include the messages we passed to debug() . The
    print(factorial(5)) call is part of the original program, so the result is dis-
    played even if logging messages are disabled.'''

def f153():
    # Logging to a file
    # Instead of displaying the log messages to the screen, you can write them to
    # a text file.

    import logging, os
    logging.basicConfig(filename='myProgramLog.txt', level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

    def factorial(n):
        logging.debug('Start of factorial(%s%%)' % (n))
        total = 1
        for i in range(n + 1):
            total *= i
            logging.debug('i is ' + str(i) + ', total is ' + str(total))
        logging.debug('End of factorial(%s%%)' % (n))
        return total

    print(factorial(5))
    os.system('cat myProgramLog.txt')
    os.unlink('myProgramLog.txt')

def f154():
    # IDLE Debuger
    print('Should be installed and practised!')

def f155():
    # Find the program bug!

    '''
    import random

    guess = ''
    while guess not in ('heads', 'tails'):
        print('Guess the coin toss! Enter heads or tails:')
        guess = input()
    toss = random.randint(0, 1) # 0 is tails, 1 is heads
    if toss == guess:
        print('You got it!')
    else:
        print('Nope! Guess again!')
        guesss = input()
        if toss == guess:
            print('You got it!')
        else:
            print('Nope. You are really bad at this game.')

    '''
    print()

def f156():
    # Chapter 11 - Web Scraping
    print('''Chapter 11 - Web Scraping:
    156 -  Launch a new browser
    157 -  Open google map address from terminal or clipboard
    158 -  Download Web Page with requests.get() Function
    159 -  Checking Downloading A Web Page
    160 -  Saving Downloaded Files to the Hard Drive
    161 -  Parsing HTML with the BeautifulSoup Module
    162 -  Creating a BeautifulSoup Object from local HTML file
    163 -  Finding an Element with the select() Method
    164 -  Getting Data from an Element’s Attributes
    165 -  I’m Feeling Lucky - Google Search
    166 -  Downloading All XKCD Comics
    167 -  Starting a Selenium-Controlled Browser
    168 -  Finding Elements on the Page
    169 -  WebElement Attributes and Methods
    170 -  Clicking the Page
    171 -  Filling Out and Submitting Forms
    172 -  Sending Special Keys
    ''')

def f157():
    # Launch a new browser by webbrowser module

    import webbrowser
    webbrowser.open('https://www.google.com')

def f158():
    # Open google map address from terminal or clipboard

    import webbrowser, sys, pyperclip

    # We chop off the first element of the array(i.e. script name)
    if len(sys.argv) > 1:
        # Get address from command line.
        address = ' '.join(sys.argv[1:])
        # address variable will be a string of commmand line arguments,
        # which are separated with space, using list join method
    else:
        # Get address from clipboard.
        address = pyperclip.paste()

    webbrowser.open('https://www.google.com/maps/place/' + address)
    # The address variable is the address in clipboard, or typed in
    # clipboard, in a browser, you can open the google map url and
    # append your address. Websites often add extra data to URLs to
    # help to track visitors or customize sites.

def f159():
    # Downloading a Web Page with the requests.get() Function

    # First install requests module:
    # sudo pip3 install requests

    import requests

    res = requests.get('http://www.gutenberg.org/cache/epub/1112/pg1112.txt')
    print(type(res))   # returns a Response object, which contains the
                        # response that the web server gave for your request.
    if res.status_code == requests.codes.ok:
        print('Request for url succeeded!')

    # If the request succeeded, the downloaded web page is stored as a string
    # in the Response object’s text variable.

    print(len(res.text)) # shows nr of characters for downloaded text
    print(res.text[:250])  # displays only the first 250 characters

def f160():
    # Checking Downloading A Web Page

    import requests

    res = requests.get('http://inventwithpython.com/page_that_does_not_exist')

    # Raise the error when the request fails and by
    # using try/except, avoid crashing the program:

    try:
        res.raise_for_status()
    except Exception as exc:
        print('There was a problem: %s' % (exc))

    # Always call raise_for_status() after calling requests.get() . You want to be
    # sure that the download has actually worked before your program continues.

def f161():
    # Saving Downloaded Files to the Hard Drive


    import requests, os

    # 1.
    # Create a Response object from URL address
    res = requests.get('http://www.gutenberg.org/cache/epub/1112/pg1112.txt')
    res.raise_for_status()

    # 2.
    # open the file in write binary mode (wb), even if the page is in
    # plain-text, to maintain the Unicode encod-ing of the text.
    playFile = open('RomeoAndJuliet.txt', 'wb')

    # 3.
    # Loop over the Response object’s iter_content() method.
    # This avoids reading the content at once into memory for large responses.
    for chunk in res.iter_content(100000): # chunk size is 100000 bytes
        # 4.
        # Call write() on each iteration to write the content to the file.
        playFile.write(chunk)

    # 5.
    # Call close() to close the file.
    playFile.close()
    print('----------')
    os.system('ls -l RomeoAndJuliet.txt')
    os.unlink('RomeoAndJuliet.txt')

def f162():
    # Parsing HTML with the BeautifulSoup Module
    import requests, bs4

    res = requests.get('http://nostarch.com')
    res.raise_for_status()
    noStarchSoup = bs4.BeautifulSoup(res.text, features="html.parser")
    print(type(noStarchSoup))

    '''This code uses requests.get() to download the main page from the
    the website and then passes the text attribute of the response to
    bs4.BeautifulSoup() . The BeautifulSoup object that it returns is
    stored in a variable named noStarchSoup .
    Once you have a BeautifulSoup object, you can use its methods to
    locate specific parts of an HTML document.'''

def f163():
    # Creating a BeautifulSoup Object from a local HTM
    import requests, bs4, os

    htmltext = '''<!-- This is the example.html example file. -->
    <html><head><title>The Website Title</title></head>
    <body>
    <p>Download my <strong>Python</strong> book from <a href="http://
    inventwithpython.com">my website</a>.</p>
    <p class="slogan">Learn Python the easy way!</p>
    <p>By <span id="author">Al Sweigart</span></p>
    </body></html>'''

    myFile = open('example.html', 'w')
    myFile.write(htmltext)
    myFile.close()

    exampleFile = open('example.html')
    exampleSoup = bs4.BeautifulSoup(exampleFile, features="html.parser")
    print(type(exampleSoup))

    os.system('rm -f example.html')

def f164():
    # Finding an Element with the select() Method

    import bs4, os

    htmltext = '''<!-- This is the example.html example file. -->
    <html><head><title>The Website Title</title></head>
    <body>
    <p>Download my <strong>Python</strong> book from <a href="http://
    inventwithpython.com">my website</a>.</p>
    <p class="slogan">Learn Python the easy way!</p>
    <p>By <span id="author">Al Sweigart</span></p>
    </body></html>'''

    myFile = open('example.html', 'w')
    myFile.write(htmltext)
    myFile.close()

    exampleFile = open('example.html')
    exampleSoup = bs4.BeautifulSoup(exampleFile.read(), features="html.parser")

    # The following matches any element that has an id attribute of author,
    # as long as it is also inside a <p> element.
    elems = exampleSoup.select('#author')    # select stors a list of Tag
                                                # objects in the elems variable
    print(type(elems))
    print(len(elems))          # tells us there is one Tag object in the list
    print(type(elems[0]))
    print(elems[0].getText())  # returns the element’s text, or inner HTML
    print(str(elems[0]))
    print(elems[0].attrs)

    print('-----------')
    pElems = exampleSoup.select('p')  # pull all the <p> elements from the
                                        # BeautifulSoup object, it gives 3 matches
    print(str(pElems[0]))
    print(pElems[0].getText())

    print(str(pElems[1]))
    print(pElems[1].getText())

    print(str(pElems[2]))
    print(pElems[2].getText())

    os.system('rm -f example.html')

def f165():
    # Getting Data from an Element’s Attributes

    # The get() method for Tag objects makes it simple to access attribute values
    # from an element. The method is passed a string of an attribute name and
    # returns that attribute’s value.

    import bs4, os

    htmltext = '''<!-- This is the example.html example file. -->
    <html><head><title>The Website Title</title></head>
    <body>
    <p>Download my <strong>Python</strong> book from <a href="http://
    inventwithpython.com">my website</a>.</p>
    <p class="slogan">Learn Python the easy way!</p>
    <p>By <span id="author">Al Sweigart</span></p>
    </body></html>'''

    myFile = open('example.html', 'w')
    myFile.write(htmltext)
    myFile.close()

    soup = bs4.BeautifulSoup(open('example.html'), features="html.parser")
    spanElem = soup.select('span')[0]
    print(str(spanElem))
    print(spanElem.get('id'))     # Passing the attribute name 'id' to get()
                                    # returns the attribute’s value, 'author'
    print(spanElem.get('some_nonexistent_addr') == None)
    print(spanElem.attrs)

    os.system('rm -f example.html')

def f166():
    # I’m Feeling Lucky - Google Search

    import requests, sys, webbrowser, bs4

    ### Step 1: Get the Command Line Arguments and Request the Search Page
    # (Instead we get user input as search pattern)
    # if you want to have the search terms as arguments of a separate running
    # script, then use the following line of code:
    # res = requests.get('http://google.com/search?q=' + ' '.join(sys.argv[1:]))

    pattern = input('Enter your search pattern: ')
    pattern = "' {} '".format(pattern)   # enclose variable with quotes
    res = requests.get('http://google.com/search?q=' +  pattern )

    print('Googling...')  # display text while downloading the Google page
    res.raise_for_status()

    ### Step 2: Find All the Results
    # Retrieve top search result links:
    soup = bs4.BeautifulSoup(res.text, features="html.parser")
    linkElems = soup.select('.r a') # returns a list of all css class type 'r',
                                    # which are in link 'a'

    ### Step 3: Open Web Browsers for Each Result
    # number of tabs you want to open is either 5 or the length of this list:
    numOpen = min(5, len(linkElems))
    for i in range(numOpen):            # Open a browser tab for each result.
        webbrowser.open('http://google.com' + linkElems[i].get('href'))

def f167():
# Downloading All XKCD Comics
    import requests, os, bs4

    ### Step 1 - TODO: Download the web page.
    url = 'http://xkcd.com'             # starting url
    os.makedirs('xkcd', exist_ok=True)  # store comics in ./xkcd

    while not url.endswith('#'):
        print('Downloading page %s...' % url)
        res = requests.get(url)
        res.raise_for_status()

        soup = bs4.BeautifulSoup(res.text, features="html.parser")

    ### Step 2 - TODO: Find the URL of the comic image.
        comicElem = soup.select('#comic img')
        if comicElem == []:
            print('Could not find comic image.')
        else:
            comicUrl = comicElem[0].get('src').strip("http://")
            comicUrl="http://"+comicUrl
            # Download the image.
            print('Downloading image %s...' % (comicUrl))
            res = requests.get(comicUrl)
            res.raise_for_status()

    ### Step 3: Save the Image and Find the Previous Comic
    # 100000 bytes of image size will be written each time to the file
    # join funcion puts an slash after folder name: xkcd/file_name.png
        imageFile = open(os.path.join('xkcd', os.path.basename(comicUrl)), 'wb')
        for chunk in res.iter_content(100000):
            imageFile.write(chunk)
        imageFile.close()

    ### Step 4: Get the Prev button's url, and repeating the loop
        prevLink = soup.select('a[rel="prev"]')[0]
        url = 'http://xkcd.com' + prevLink.get('href')
        # the selector 'a[rel="prev"]' identifies the <a> element with
        # the rel attribute set to prev , and you can use this <a> element’s href
        # attribute to get the previous comic’s URL, which gets stored in url.
        # Now the same process continue for the new url, up to when url ends with #.

    print('Done.')

def f168():
    # Starting a Selenium-Controlled Browser

    from selenium import webdriver
    browser = webdriver.Firefox()     # opens firefox browser, with no url loaded
    print(type(browser))              # prints the class type of browser:
    # <class 'selenium.webdriver.firefox.webdriver.WebDriver'>
    browser.get('http://inventwithpython.com')  # opens the url in the already opened
                                                # firefox browser

def f169():
    # Finding Elements on the Page

    from selenium import webdriver
    browser = webdriver.Firefox()  # opens firefox browser, with no url loaded
    browser.get('http://inventwithpython.com') # loads the url argument in browser

    try:
        elem = browser.find_element_by_class_name('card-img-top')
        # The find_element_* methods is looking just for a single element,
        # while find_elements_* methods returns a list of WebElement_* objects.
        print('Found <%s> element with that class name!' % (elem.tag_name))
    except:
        print('Was not able to find an element with that name.')

    # The output will be:
    # Found <img> element with that class name!

    # To see the list of methods refer to item 159 in main menu !

def f170():
    # WebElement Attributes and Methods

    ''' Selenium’s WebDriver Methods for Finding Elements
    - Elements that use the CSS class name:
    browser.find_element_by_class_name(name)
    browser.find_elements_by_class_name(name)

    - Elements that match the CSS selector
    browser.find_element_by_css_selector(selector)
    browser.find_elements_by_css_selector(selector)

    - Elements with a matching id attribute value
    browser.find_element_by_id(id)
    browser.find_elements_by_id(id)

    - <a> elements that completely match the text provided
    browser.find_element_by_link_text(text)
    browser.find_elements_by_link_text(text)

    - <a> elements that contain the text provided
    browser.find_element_by_partial_link_text(text)
    browser.find_elements_by_partial_link_text(text)

    - Elements with a matching name attribute value
    browser.find_element_by_name(name)
    browser.find_elements_by_name(name)

    - Elements with a matching tag name
      (case insensitive; an <a> element is matched by 'a' and 'A')
    browser.find_element_by_tag_name(name)
    browser.find_elements_by_tag_name(name)

    tag_name - The tag name, such as 'a' for an <a> element
    get_attribute(name) - The value for the element’s name attribute
    text - The text within the element, such as 'hello' in <span>hello</span>
    clear() - For text field or text area elements, clears the text typed into it
    is_displayed() - Returns True if the element is visible; otherwise returns False
    is_enabled() - For input elements, returns True if the element is enabled; other-
                   wise returns False
    is_selected() - For checkbox or radio button elements, returns True if the ele-
                    ment is selected; otherwise returns False
    location - A dictionary with keys 'x' and 'y' for the position of the ele-
               ment in the page
    '''

def f171():
    # Clicking the Page

    from selenium import webdriver
    browser = webdriver.Firefox()   # opens firefox browser, with no url loaded
    browser.get('http://inventwithpython.com')   # loads the url argument in browser
    linkElem = browser.find_element_by_link_text('Read Online for Free')
    print(type(linkElem))
    # It prints: <class 'selenium.webdriver.remote.webelement.WebElement'>
    linkElem.click()    # follows the "Read It Online" link
                        # click() simulate mouse click on the first link element with
                        # text as 'Read Online for Free'

def f172():
    # Filling Out and Submitting Forms

    from selenium import webdriver
    browser = webdriver.Firefox()   # opens firefox browser, with no url loaded
    browser.get('http://gmail.com') # loads the url argument in browser
    emailElem = browser.find_element_by_id('identifierId')
    emailElem.send_keys('not_my_real_email@gmail.com')
    passwordElem = browser.find_element_by_id('Passwd')
    passwordElem.send_keys('12345')
    passwordElem.submit()

    '''The previous code will fill in those text fields with the provided text.
    (You can always use the browser’s inspector to verify the id .) Calling the
    submit() method on any element will have the same result as clicking the Submit
    button for the form that element is in.'''

def f173():
    # Sending Special Keys

    from selenium import webdriver
    from selenium.webdriver.common.keys import Keys

    browser = webdriver.Firefox() # opens just the firefox browser,
                                    # no url loaded is loaded
    browser.get('http://nostarch.com') # loads url in browser
    htmlElem = browser.find_element_by_tag_name('html')
    htmlElem.send_keys(Keys.END)   # scrolls to bottom
    # To scroll up again uncomment the following
    # htmlElem.send_keys(Keys.HOME)  # scrolls to top

    '''Commonly Used Variables in the selenium.webdriver.common.keys Module
    Keys.DOWN , Keys.UP , Keys.LEFT , Keys.RIGHT     The keyboard arrow keys
    Keys.ENTER , Keys.RETURN                         The Enter or RETURNS keys
    Keys.HOME , Keys.END , Keys.PAGE_DOWN, Keys.PAGE_UP   HOME,END, ...
    Keys.ESCAPE , Keys.BACK_SPACE , Keys.DELETE       esc , backspace , DELETE
    Keys.F1 , Keys.F2 , . . . , Keys.F12              F1 to F12
    Keys.TAB                                          The tab key

    - Selenium can simulate clicks on various browser buttons by
      following methods:

    browser.back()  Clicks the Back button.
    browser.forward()  Clicks the Forward button.
    browser.refresh()  Clicks the Refresh/Reload button.
    browser.quit()  Clicks the Close Window button.

    Selenium reference: http://selenium-python.readthedocs.org/
    '''
def f174():
    #
    import socket

    mysock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    mysock.connect(('data.pr4e.org', 80))
    cmd = 'GET http://data.pr4e.org/romeo.txt HTTP/1.0\r\n\r\n'.encode()
    mysock.send(cmd)

    while True:
        data = mysock.recv(512)
        if len(data) < 1:
            break
        print(data.decode(),end='')

    mysock.close()

def f175():
    # Download an image from webserver while showing the file
    # header response information and dowloaded chuncks of data,
    # during the download

    import socket, time

    HOST = 'data.pr4e.org'
    PORT = 80
    mysock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    mysock.connect((HOST, PORT))
    mysock.sendall(b'GET http://data.pr4e.org/cover3.jpg HTTP/1.0\r\n\r\n')
    count = 0
    picture = b""

    while True:
        data = mysock.recv(5120)
        if len(data) < 1: break
        time.sleep(0.25)
        count = count + len(data)
        print(len(data), count)
        picture = picture + data

    mysock.close()

    # Look for the end of the header (2 CRLF)
    pos = picture.find(b"\r\n\r\n")
    print('Header length', pos)
    print(picture[:pos].decode())

    # Skip past the header and save the picture data
    picture = picture[pos+4:]
    fhand = open("stuff.jpg", "wb")
    fhand.write(picture)
    fhand.close()

def f176():
    # Retrieve a web page over a socket and display the headers
    # from the web server.

    import socket

    HOST = input('Enter webserver URL: ')
    DOC = input('Enter document to be retrived [Press Enter to default index.html]: ')
    if DOC == '': DOC = 'index.html'
    PORT = input('Enter port number [Press Enter to default 80]: ')
    if PORT == '': PORT = 80

    # Open a connection
    mysock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    mysock.connect((HOST, PORT))


    cmd = 'GET http://' +  HOST +  '/' + DOC + ' HTTP/1.0\r\n\r\n'
    cmd = cmd.encode()
    print(cmd)

    mysock.send(cmd)


    receivedData = b""

    while True:
        data = mysock.recv(512)
        if len(data) < 1: break
        receivedData = receivedData + data

    mysock.close()

    pos = receivedData.find(b"\r\n\r\n")
    print('Header length', pos)
    print(receivedData[:pos].decode())

def f177():
    # Load a web document, browser simulation
    import socket

    mysock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    mysock.connect(('data.pr4e.org', 80))
    cmd = 'GET http://data.pr4e.org/intro-short.txt HTTP/1.0\r\n\r\n'.encode()
    mysock.send(cmd)

    while True:
        data = mysock.recv(512)
        if len(data) < 1:
            break
        print(data.decode(),end='')

    mysock.close()

def f178():
    # Open a web page with urlib module
    import urllib.request

    fhand = urllib.request.urlopen('http://data.pr4e.org/romeo.txt')
    for line in fhand:
        print(line.decode().strip())

    # This program displays the html document as you see the sour-
    # ce code in browser. The urlopen returns an object similar to
    # file handle that can be  manipulated the same way as a file.
    # For example, you can iterate each line of  html document, or
    # any other oprations.

def f179():
    # Count the number of words in a web page
    import urllib.request, urllib.parse, urllib.error, pprint

    fhand = urllib.request.urlopen('http://data.pr4e.org/romeo.txt')

    counts = dict()   # Creates an empty dictionary
    for line in fhand:
        words = line.decode().split() # line coming from webserver
                                      # will be decoded and split-
                                      # ted as words.
                                      # decode converts to string.

        for word in words:   # Use a dictionary to count the words
            counts[word] = counts.get(word, 0) + 1
            # get(dic_key, set_default) is used to get value  for a
            # key in dictionay, if not exists, here is set to 0.

    pprint.pprint(counts)

    # Here is another operation, counting number of the words in a
    # html document. Here html document, has been  treated  like a
    # regular file, with urlopen() function.

def f180():
    # Retrieve href links from html document with BeautifulSoup
    # To run this, you can install BeautifulSoup
    # https://pypi.python.org/pypi/beautifulsoup4

    # Or download the file
    # http://www.py4e.com/code3/bs4.zip
    # and unzip it in the same directory as this file

    import urllib.request, urllib.parse, urllib.error
    from bs4 import BeautifulSoup
    import ssl

    # Ignore SSL certificate errors
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE

    url = input('Enter - ')
    html = urllib.request.urlopen(url, context=ctx).read()
    soup = BeautifulSoup(html, 'html.parser')

    # Retrieve all of the anchor tags
    tags = soup('a')
    for tag in tags:
        print(tag.get('href', None))

def f181():
    # Chapter 12 - Working with Excel Spreadsheets
    print('Chapter 12 - Working with Excel Spreadsheets')

def f182():
    # Opening Excel Documents with OpenPyXL
    # To work with Excel documents first install the module:
    # sudo pip3 install OpenPyXL

    import openpyxl
    wb = openpyxl.load_workbook('example.xlsx')
    # This does not open the file

    print(type(wb))  # Returns workbook data type

    # printing list of all the sheet names in the workbook
    print(wb.get_sheet_names())
    sheet = wb.get_sheet_by_name('Sheet3')
    # Returns only the name of Sheet3 to variable sheet
    print(sheet)

    print(type(sheet)) # Returns Worksheet data type
    print(sheet.title) #

    # The active sheet is the sheet that’s on top when the
    # workbook is opened in Excel.

    # To get the workbook’s active sheet:
    anotherSheet = wb.get_active_sheet()
    print(anotherSheet)

def f183():
    # Getting Cells from the Sheets

    import openpyxl
    wb = openpyxl.load_workbook('example.xlsx')
    # This does not open the file, just loads the file in memory

    sheet = wb.get_sheet_by_name('Sheet1')
    print(sheet['A1'])
    print(sheet['A1'].value) # Note: cell A1 is quoted!
    c = sheet['B1'] # Note: cell B1 is quoted!
    print(c.value)

    print('Row ' + str(c.row) + ', Column ' + str(c.column) + ' is '
          + c.value)
    print('Cell ' + c.coordinate + ' is ' + c.value)
    print(sheet['C1'].value)

    # 'value' is an attribute of Cell objects.
    # Cell objects also have row , column , and coordinate attri-
    # butes that provide location information for the cell.

def f184():
    # excel sheet cell() method

    import openpyxl
    wb = openpyxl.load_workbook('example.xlsx')
    sheet = wb.get_sheet_by_name('Sheet1')

    sheet.cell(row=1, column=2)         # <Cell Sheet1.B1>
    sheet.cell(row=1, column=2).value   # 'Apples'
    for i in range(1, 8, 2):
        print(i, sheet.cell(row=i, column=2).value)

def f185():
    # Determine the size of the sheet

    import openpyxl
    wb = openpyxl.load_workbook('example.xlsx')
    sheet = wb.get_sheet_by_name('Sheet1')

    # Following deprecated:
    # print(sheet.get_highest_row())                 # 7
    # print(sheet.get_highest_column())              # 3

    print(sheet.max_row)                 # 7
    print(sheet.max_column)              # 3

def f186():
    # Converting Between Column Letters and Numbers
    import openpyxl
    from openpyxl.utils import get_column_letter, column_index_from_string

    get_column_letter(1)         # 'A'
    get_column_letter(2)         # 'B'
    get_column_letter(27)        # 'AA'
    get_column_letter(900)       # 'AHP'

    wb = openpyxl.load_workbook('example.xlsx')
    sheet = wb.get_sheet_by_name('Sheet1')

    # Following deprecated:
    # get_column_letter(sheet.get_highest_column())   #'C'
    print(get_column_letter(sheet.max_column))        #'C'
    print(column_index_from_string('A'))              # 1
    print(column_index_from_string('AA'))             # 27

def f187():
    # Getting Rows and Columns from the Sheets

    import openpyxl

    wb = openpyxl.load_workbook('example.xlsx')
    sheet = wb.get_sheet_by_name('Sheet1')
    print(tuple(sheet['A1':'C3']))

    for rowOfCellObjects in sheet['A1':'C3']:
        for cellObj in rowOfCellObjects:
            print(cellObj.coordinate, cellObj.value)
        print('--- END OF ROW ---')

    # To print the values of each cell in the area, we use two for
    # loops. The outer for loop goes over each row in the slice.
    # Then, for each row, the nested for loop goes through each
    # cell in that row.

def f188():
    # Access the values of cells in a particular row or column

    import openpyxl

    wb = openpyxl.load_workbook('example.xlsx')
    sheet = wb.get_active_sheet()
    print(list(sheet.columns)[1])

    for cellObj in list(sheet.columns)[1]:
        print(cellObj.value)

    #    Quick Review
    # 1. Import the openpyxl module.
    # 2. Call the openpyxl.load_workbook() function.
    # 3. Get a Workbook object.
    # 4. Call the get_active_sheet() or get_sheet_by_name()
    #    workbook method.
    # 5. Get a Worksheet object.
    # 6. Use indexing or the cell() sheet method with row and
    #    column keyword arguments.
    # 7. Get a Cell object.
    # 8. Read the Cell object’s value attribute.

def f189():
    # Reading Data from a Spreadsheet
    # Tabulates population and number of census tracts for
    # each county.

    import openpyxl, pprint
    print('Opening workbook...')
    wb = openpyxl.load_workbook('censuspopdata.xlsx')
    sheet = wb.get_sheet_by_name('Population by Census Tract')
    countyData = {}

    # Fill in countyData with each county's population and tracts.
    print('Reading rows...')
    # iterating over the rowa

    # method sheet.get_highest_row(), has been deprecated and ins-
    # tead we use sheet.max_row

    for row in range(2, sheet.max_row + 1):
        # Each row in the spreadsheet has data for one census tract.
        state  = sheet['B' + str(row)].value
        county = sheet['C' + str(row)].value
        pop    = sheet['D' + str(row)].value

        # Make sure the key for this state exists.
        countyData.setdefault(state, {})
        # Make sure the key for this county in this state exists.
        countyData[state].setdefault(county, {'tracts': 0, 'pop': 0})

        # Each row represents one census tract, so increment by one.
        countyData[state][county]['tracts'] += 1
        # Increase the county pop by the pop in this census tract.
        countyData[state][county]['pop'] += int(pop)

    # Open a new text file and write the contents of countyData to it.
    print('Writing results...')
    resultFile = open('census2010.py', 'w')
    resultFile.write('allData = ' + pprint.pformat(countyData))
    resultFile.close()
    print('Done.')

def f190():
    # Writing Excel Documents
    import openpyxl

    wb = openpyxl.Workbook() #creates a new, blank Workbook object
    # No document wil be saved unil usinfg save method.

    wb.get_sheet_names()
    sheet = wb.get_active_sheet()
    print(sheet.title)
    sheet.title = 'Spam Bacon Eggs Sheet'
    wb.get_sheet_names()

def f191():
    # Create and save a spreed sheet
    import openpyxl

    wb = openpyxl.load_workbook('example.xlsx')
    sheet = wb.get_active_sheet()
    sheet.title = 'Spam Spam Spam'
    wb.save('example_copy.xlsx')

def f192():
    # Creating and Removing Sheets

    import openpyxl

    wb = openpyxl.Workbook()
    print(wb.get_sheet_names())

    wb.create_sheet(index=0, title='First Sheet')
    print(wb.get_sheet_names())

    wb.create_sheet(index=2, title='Middle Sheet')
    print(wb.get_sheet_names())


    wb.remove_sheet(wb.get_sheet_by_name('Middle Sheet'))
    print(wb.get_sheet_names())
    wb.remove_sheet(wb.get_sheet_by_name('First Sheet'))
    print(wb.get_sheet_names())

    # If you want the new created excel be saved, uncomment this:
    # wb.save('test.xlsx')

def f193():
    # Writing Values to Cells

    import openpyxl

    # Writing values to cells is much like writing values to keys
    # in a dictionary.
    wb = openpyxl.Workbook()
    sheet = wb.get_sheet_by_name('Sheet')
    sheet['A1'] = 'Hello world!'
    print(sheet['A1'].value)

    # If you want the new created excel be saved, uncomment this:
    # wb.save('test.xlsx')

def f194():
    # Project: Updating a Spreadsheet
    # This program updates prices for specific produces in file
    # 'produceSales.xlsx'

    import openpyxl
    wb = openpyxl.load_workbook('produceSales.xlsx')
    sheet = wb.get_sheet_by_name('Sheet')
    # The produce types and their updated prices
    PRICE_UPDATES = {'Garlic': 3.07,
                     'Celery': 1.19,
                    'Lemon': 1.27}

    # Loop through the rows and update the prices.
    # skip the first row:
    for rowNum in range(2, sheet.max_row):
        # The cell in column 1 (that is, column A) will be stored
        # in the ­variable produceName
        produceName = sheet.cell(row=rowNum, column=1).value
        # If produceName exists as a key in the PRICE_UPDATES
        # dictionary, then update the cell
        if produceName in PRICE_UPDATES:
            sheet.cell(row=rowNum, column=2).value = PRICE_UPDATES[produceName]

    wb.save('updatedProduceSales.xlsx')

def f195():
    # Setting the Font Style of Cells

    import openpyxl
    from openpyxl.styles import NamedStyle, Font, Border, Side

    wb = openpyxl.Workbook()
    sheet = wb.get_sheet_by_name('Sheet')
    highlight = NamedStyle(name="highlight")
    highlight.font = Font(size=24, italic=True)
    # Once a named style has been created, it can be registered
    # with the workbook:
    wb.add_named_style(highlight)
    # But named styles will also be registered automatically the
    # first time they are assigned to a cell:
    sheet['A1'].style = highlight
    sheet['A1'] = 'Hello world!'
    wb.save('styled.xlsx')

def f196():
    # Font object

    import openpyxl
    from openpyxl.styles import Font, NamedStyle

    wb = openpyxl.Workbook()
    sheet = wb.get_sheet_by_name('Sheet')

    fontObj1 = NamedStyle(name="fontObj1")
    fontObj1.font = Font(name='Times New Roman', bold=True)
    sheet['A1'].style = fontObj1
    sheet['A1'] = 'Bold Times New Roman'

    fontObj2 = NamedStyle(name="fontObj2")
    fontObj2.font = Font(size=24, italic=True)
    sheet['B3'].style = fontObj2
    sheet['B3'] = 'Bold Times New Roman'

    wb.save('styles.xlsx')

def f197():
    # Formulas
    import openpyxl

    wb = openpyxl.Workbook()
    sheet = wb.get_active_sheet()
    sheet['A1'] = 200
    sheet['A2'] = 300
    sheet['A3'] = '=SUM(A1:A2)'
    wb.save('writeFormula.xlsx')

def f198():
    # loading a workbook with and without the data_only keyword
    # argument

    import openpyxl

    wbFormulas = openpyxl.load_workbook('writeFormula.xlsx')
    sheet = wbFormulas.get_active_sheet()
    print(sheet['A3'].value)

    wbDataOnly = openpyxl.load_workbook('writeFormula.xlsx',
                                        data_only=True)
    ws = wbDataOnly.get_active_sheet()
    print(ws['A3'].value)
    # in my case it shows None

def f199():
    # Setting Row Height and Column Width

    import openpyxl

    wb = openpyxl.Workbook()
    sheet = wb.get_active_sheet()
    sheet['A1'] = 'Tall row'
    sheet['B2'] = 'Wide column'
    sheet.row_dimensions[1].height = 70
    sheet.column_dimensions['B'].width = 20
    wb.save('dimensions.xlsx')

def f200():
    # Merging and Unmerging Cells
    # A rectangular area of cells can be merged into a single cell
    # with the merge_cells() sheet method

    import openpyxl

    wb = openpyxl.Workbook()
    sheet = wb.get_active_sheet()
    sheet.merge_cells('A1:D3')
    sheet['A1'] = 'Twelve cells merged together.'
    sheet.merge_cells('C5:D5')
    sheet['C5'] = 'Two merged cells.'
    wb.save('merged.xlsx')

def f201():
    # Unmerge cells
    # To unmerge cells, call the unmerge_cells() sheet method

    import openpyxl

    wb = openpyxl.load_workbook('merged.xlsx')
    sheet = wb.get_active_sheet()
    sheet.unmerge_cells('A1:D3')
    sheet.unmerge_cells('C5:D5')
    wb.save('merged.xlsx')

def f202():
    # Freeze Panes - Pan is visible top rows or leftmost columns
    # to a specific cell that is helpfull to be used for larg spr-
    # eadsheets to display all at once.

    # sheet.freeze_panes = 'A2'    Row 1 is frozen
    # sheet.freeze_panes = 'B1'    Column A is frozen
    # sheet.freeze_panes = 'C1'    Columns A and B are frozen
    # sheet.freeze_panes = 'C2'    Row 1 and columns A and B
    # sheet.freeze_panes = 'A1'    No frozen panes
    # sheet.freeze_panes = None    No frozen panes

    import openpyxl
    wb = openpyxl.load_workbook('produceSales.xlsx')
    sheet = wb.get_active_sheet()
    sheet.freeze_panes = 'A2' # The row header si visible, even when
                                # scroll down
    wb.save('freezeExample.xlsx')
    # To unfreez all pan set the the atribute to None

def f203():
    # Charts
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    for i in range(10):
        ws.append([i])

    from openpyxl.chart import BarChart, Reference, Series
    values = Reference(ws, min_col=1, min_row=1, max_col=1, max_row=10)
    chart = BarChart()
    chart.add_data(values)
    ws.add_chart(chart, "E15")
    wb.save("sampleChart.xlsx")

    # For complete documentation refere to:
    # https://openpyxl.readthedocs.io/en/stable/charts/introduction.html

def f204():
    # Chapter 13 - Working with PDF and Word Documents
    print('Chapter 13 - Working with PDF and Word Documents')

def f205():
    # Extracting Text from PDFs
    # sudo pip3 install PyPDF2

    import PyPDF2

    pdfFileObj = open('meetingminutes.pdf', 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
    print(pdfReader.numPages)
    pageObj = pdfReader.getPage(0)
    print(pageObj.extractText())

def f206():
    # Decrypting PDFs

    import PyPDF2

    pdfReader = PyPDF2.PdfFileReader(open('encrypted.pdf', 'rb'))
    print(pdfReader.isEncrypted)
    # When you try uncomment following, you get error(encrypted)
    # pdfReader.getPage(0)
    pdfReader.decrypt('rosebud')   # password='rosebud'
    pageObj = pdfReader.getPage(0) # No more error you get
    print(pageObj.extractText())

def f207():
    # copy pages from one PDF document to another
    # files 'meetingminutes.pdf' and 'meetingminutes2.pdf' must
    # exist in current directory!

    import PyPDF2, os

    if not os.path.isfile('meetingminutes.pdf'):
        print('meetingminutes.pdf is required in current working directory!')
        return

    pdf1File = open('meetingminutes.pdf', 'rb')
    pdf2File = open('meetingminutes2.pdf', 'rb')
    pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
    pdf2Reader = PyPDF2.PdfFileReader(pdf2File)
    pdfWriter = PyPDF2.PdfFileWriter()

    for pageNum in range(pdf1Reader.numPages):
        pageObj = pdf1Reader.getPage(pageNum)
        pdfWriter.addPage(pageObj)

    for pageNum in range(pdf2Reader.numPages):
        pageObj = pdf2Reader.getPage(pageNum)
        pdfWriter.addPage(pageObj)

    pdfOutputFile = open('combinedminutes.pdf', 'wb')
    pdfWriter.write(pdfOutputFile)
    pdfOutputFile.close()
    pdf1File.close()
    pdf2File.close()


def f208():
    # Rotating Pages
    # file 'meetingminutes.pdf' must be in current directory!

    import PyPDF2, os

    if not os.path.isfile('meetingminutes.pdf'):
        print('meetingminutes.pdf is required in current working directory!')
        return

    minutesFile = open('meetingminutes.pdf', 'rb')
    pdfReader = PyPDF2.PdfFileReader(minutesFile)
    page = pdfReader.getPage(0)
    page.rotateClockwise(90)

    pdfWriter = PyPDF2.PdfFileWriter()
    pdfWriter.addPage(page)
    resultPdfFile = open('rotatedPage.pdf', 'wb')
    pdfWriter.write(resultPdfFile)
    resultPdfFile.close()
    minutesFile.close()

def f209():
    # Overlaying Pages
    # files 'meetingminutes.pdf' and 'watermark.pdf' must exist in
    # current directory!
    import os

    if not os.path.isfile('meetingminutes.pdf') and os.path.isfile('watermark.pdf'):
        print('meetingminutes.pdf and watermark.pdf are required!')
        return

    import PyPDF2

    minutesFile = open('meetingminutes.pdf', 'rb')
    pdfReader = PyPDF2.PdfFileReader(minutesFile)
    minutesFirstPage = pdfReader.getPage(0)
    pdfWatermarkReader = PyPDF2.PdfFileReader(open('watermark.pdf', 'rb'))
    minutesFirstPage.mergePage(pdfWatermarkReader.getPage(0))
    pdfWriter = PyPDF2.PdfFileWriter()
    pdfWriter.addPage(minutesFirstPage)

    for pageNum in range(1, pdfReader.numPages):
        pageObj = pdfReader.getPage(pageNum)
        pdfWriter.addPage(pageObj)

    resultPdfFile = open('watermarkedCover.pdf', 'wb')
    pdfWriter.write(resultPdfFile)
    minutesFile.close()
    resultPdfFile.close()

def f210():
    # Encrypting PDFs
    import PyPDF2
    pdfFile = open('meetingminutes.pdf', 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFile)
    pdfWriter = PyPDF2.PdfFileWriter()
    for pageNum in range(pdfReader.numPages):
        pdfWriter.addPage(pdfReader.getPage(pageNum))

    pdfWriter.encrypt('swordfish')
    resultPdf = open('encryptedminutes.pdf', 'wb')
    pdfWriter.write(resultPdf)
    resultPdf.close()

def f211():
    # Combining Select Pages from Many PDFs

    import PyPDF2, os

    # Get all the PDF filenames.
    pdfFiles = []
    for filename in os.listdir('.'):
        if filename.endswith('.pdf'):
            pdfFiles.append(filename)

    pdfFiles.sort(key=str.lower)
    pdfWriter = PyPDF2.PdfFileWriter()

    # Loop through all the PDF files.
    for filename in pdfFiles:
        pdfFileObj = open(filename, 'rb')
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
        # TODO: Loop through all the pages (except the first) and
        # add them.
        for pageNum in range(1, pdfReader.numPages):
            pageObj = pdfReader.getPage(pageNum)
            pdfWriter.addPage(pageObj)

    # Save the resulting PDF to a file.
    pdfOutput = open('allminutes.pdf', 'wb')
    pdfWriter.write(pdfOutput)
    pdfOutput.close()

def f212():
    # Word Documents
    # sudo pip3 install python-docx (install module 'python-docx')
    # file 'demo.docx' must be in working directory.

    import docx

    doc = docx.Document('demo.docx') # return a Document object
    print(len(doc.paragraphs)) # a list of paragraph objects

    # string of the text in paragraph[0]:
    print(doc.paragraphs[0].text)
    # string of the text in paragraph[1]
    print(doc.paragraphs[1].text)

    # Each paragraph object has a list of Run objects
    print(len(doc.paragraphs[1].runs))
    print(doc.paragraphs[1].runs[0].text)
    print(doc.paragraphs[1].runs[1].text)
    print(doc.paragraphs[1].runs[2].text)
    print(doc.paragraphs[1].runs[3].text)

def f213():
    # Getting the Full Text from a .docx File

    import docx

    def getText(filename):
        doc = docx.Document(filename)
        fullText = []
        for para in doc.paragraphs:
            # fullText.append(para.text)
            fullText.append(' ' + para.text)

        return '\n'.join(fullText)

    print(getText('demo.docx'))

def f214():
    # Styling Paragraph and Run Attributes
    print('''
    The string values for the default Word styles are:
    Normal BodyText BodyText2 BodyText3 Caption Heading1 Heading2
    Heading3 Heading4 Heading5 Heading6 Heading7 Heading8 Heading9
    IntenseQuote List List2 List3 ListBullet ListBullet2 ListBullet3
    ListContinue ListContinue2 ListContinue3 ListNumber ListNumber2
    ListNumber3 ListParagraph MacroText NoSpacing Quote Subtitle
    TOCHeading Title
    ''')

    print('''
    Run attributes are:
    bold italic underline strike double_strike all_caps small_caps
    shadow outline rtl imprint emboss
    ''')


def f215():
    # Creating Word Documents with Nondefault Styles
    import docx

    doc = docx.Document('demo.docx')
    doc.paragraphs[0].text
    'Document Title'
    print(doc.paragraphs[0].style)    #'Title'
    doc.paragraphs[0].style = 'Normal'
    print(doc.paragraphs[1].text)
    # returns:
    # 'A plain paragraph with some bold and some italic'
    (doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.
     paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)
    # output:
    # ('A plain paragraph with some ', 'bold', ' and some ', 'italic')

    doc.paragraphs[1].runs[0].style = 'QuoteChar'
    doc.paragraphs[1].runs[1].underline = True
    doc.paragraphs[1].runs[3].underline = True
    doc.save('restyled.docx')

def f216():
    # Writing Word Documents
    import docx

    doc = docx.Document()
    print(doc.add_paragraph('Hello world!'))
    doc.save('helloworld.docx')

def f217():
    # Add paragraphs

    import docx

    doc = docx.Document()
    doc.add_paragraph('Hello world!')

    paraObj1 = doc.add_paragraph('This is a second paragraph.')
    paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
    paraObj1.add_run(' This text is being added to the second paragraph.')

    doc.save('multipleParagraphs.docx')

def f218():
    # Adding Headings

    import docx

    doc = docx.Document()
    doc.add_heading('Header 0', 0)
    doc.add_heading('Header 1', 1)
    doc.add_heading('Header 2', 2)
    doc.add_heading('Header 3', 3)
    doc.add_heading('Header 4', 4)
    doc.save('headings.docx')

def f219():
    # Adding Line and Page Breaks and Adding Pictures

    import docx, os

    if not os.path.isfile('zophie.png'):
        print('zophie.png is required in current working directory!')
        return
    doc = docx.Document()
    doc.add_paragraph('This is on the first page!')
    # Following didn't work, for page break
    # doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE)
    doc.add_page_break()
    doc.add_paragraph('This is on the second page!')
    # doc.save('twoPage.docx')
    doc.add_picture('zophie.png', width=docx.shared.Inches(1),
                    height=docx.shared.Cm(4))
    doc.save('twoPage.docx')

def f220():
    # Chapter 14 - Working with CSV Files and JSON Data
    print('Chapter 14 - Working with CSV Files and JSON Data')

def f221():
    # Read data from a CSV file with csv module,
    # you need to create a Reader object

    # csv is built in module
    import csv, os
    from pprint import pprint

    if not os.path.isfile('example.csv'):
        print('example.csv is required in current working directory!')
        return
    # open a file just any other text file, with open()
    exampleFile = open('example.csv')

    # instead of read/readlines for files use reader method:
    exampleReader = csv.reader(exampleFile) # returns a Reader obj

    # To access values, convert the reader object to a list:
    exampleData = list(exampleReader)
    pprint(exampleData)

def f222():
    # Reading Data from Reader Objects in a for Loop

    import csv, os

    if not os.path.isfile('example.csv'):
        print('example.csv is required in current working directory!')
        return

    exampleFile = open('example.csv')
    exampleReader = csv.reader(exampleFile)

    for row in exampleReader:
        # to get the row number in Reader object use
        # variable line_num
        print('Row #' + str(exampleReader.line_num) + \
              ' ' + str(row))

        # The Reader object can be looped over only once

def f223():
    # Writer object
    # A Writer object lets you write data to csv file.

    import csv

    # Open a file in write mode:
    outputFile = open('output.csv', 'w', newline='')
    # If you forget the newline='' keyword argument in open(), the
    # CSV file will be double-spaced.

    # Create a Writer object:
    outputWriter = csv.writer(outputFile)

    outputWriter.writerow(['spam', 'eggs', 'bacon', 'ham'])
    outputWriter.writerow(['Hello, world!', 'eggs', 'bacon', 'ham'])
    outputWriter.writerow([1, 2, 3.141592, 4])
    outputFile.close()

    with open('output.csv') as myFile: print(myFile.read())

def f224():
    # The 'delimiter' and 'lineterminator' keyword arguments
    # We change the above arguments to change the default values.
    # Arguments default values are comma and '\n'.

    import csv

    csvFile = open('example.tsv', 'w', newline='')
    csvWriter = csv.writer(csvFile, delimiter='\t',
                           lineterminator='\n\n')

    csvWriter.writerow(['apples', 'oranges', 'grapes'])
    csvWriter.writerow(['eggs', 'bacon', 'ham'])
    csvWriter.writerow(['spam', 'spam', 'spam', 'spam', 'spam',
                        'spam'])

    csvFile.close()

def f225():
    # Project: Removes the header from all CSV files in the current
    # working directory.

    import csv, os

    os.makedirs('headerRemoved', exist_ok=True)

    # Step 1: Loop Through Each CSV File
    # ...................................
    for csvFilename in os.listdir('.'):
        if not csvFilename.endswith('.csv'):
            continue
            # makes the for loop move on to the next
            # filename when it comes across a non-CSV file

        print('Removing header from ' + csvFilename + '...')

        # Step 2: Read in the CSV File (skipping first row)
        # .................................................
        csvRows = []
        csvFileObj = open(csvFilename)
        readerObj = csv.reader(csvFileObj)
        for row in readerObj:
            if readerObj.line_num == 1:
                continue  # skip first row
            csvRows.append(row)
        csvFileObj.close()

        # Step 3: Write Out the CSV File Without the First Row
        # .....................................................
        csvFileObj = open(os.path.join('headerRemoved',
                                       csvFilename), 'w', newline='')

        csvWriter = csv.writer(csvFileObj)
        for row in csvRows:
            csvWriter.writerow(row)
        csvFileObj.close()

def f226():
    # What is JSON
    print('''
    JavaScript Object Notation is a popular way to format data as
    a single human-readable string. JSON is the native way that J-
    avaScript programs write their data structures and usually re-
    sembles what Python’s pprint function would produce. Example:
    {"name": "Zophie", "isCat": true,
    "miceCaught": 0, "napsTaken": 37.5,
    "felineIQ": null}

    Many websites offer JSON content as a way for programs to int-
    eract with the website. This is known as providing an API. Ac-
    cessing an API is the same as accessing any other web page via
    a URL. The difference is that the data returned by an API is -
    formatted with JSON. Facebook, Twitter, Yahoo, Google, Tumblr,
    Wikipedia, Flickr, Data.gov, Reddit, IMDb, Rotten Tomatoes, L-
    inkedIn, and many other offer APIs for programs to use. Some -
    of these sites require registration, which is almost always f-
    ree. You’ll have to find documentation for what URLs your pro-
    gram needs to request in order to get the data you want, as w-
    ell as the general format of the JSON data structures that are
    returned. This documentation should be provided by whatever s-
    ite is offering the API; if they have a “Developers” page, lo-
    ok for the documentation there. What you can do with JSOn:
    - Scrape raw data from websites.
      Accessing APIs and parsing HTML with Beautiful Soup.
    - Automatically download new posts from one of your social ne-
      twork accounts and post them to another account.
    - Pulling data from IMDb, Rotten Tomatoes, and Wikipedia and -
      putting it into a single text file.
    ''')

def f227():
    # Reading JSON with the loads() Function
    # To translate a string containing JSON data into a Python
    # value, pass it to the json.loads() function.

    # JSON strings always use double quotes
    stringOfJsonData = '{"name": "Zophie", "isCat": true, \
                        "miceCaught": 0, "felineIQ": null}'

    import json   # built-in module
    jsonDataAsPythonValue = json.loads(stringOfJsonData)

    print(jsonDataAsPythonValue)

def f228():
    # Writing JSON with the dumps() Function
    # The json.dumps() function ("dump string") translate a Python
    # value into a string of JSON-formatted data.

    pythonValue = {'isCat': True, 'miceCaught': 0,
                   'name': 'Zophie', 'felineIQ': None}

    import json
    stringOfJsonData = json.dumps(pythonValue)

    print(stringOfJsonData)

def f229():
    # Project: Fetching Current Weather Data

    import json, requests, sys

    # Step 1: Get Location from the Command Line Argument
    # # Compute location from command line arguments.
    # if len(sys.argv) < 2:
    #     print('Usage: quickWeather.py location')
    #     sys.exit()

    # location = ' '.join(sys.argv[1:])
    location = 'San Francisco CA'

    # Step 2: Download the JSON Data
    url ='http://api.openweathermap.org/data/2.5/forecast/daily?q=%s&cnt=3' % (location)
    response = requests.get(url)
    response.raise_for_status()

    # Load JSON data into a Python variable.
    weatherData = json.loads(response.text)

    # Print weather descriptions.
    w = weatherData['list']
    print('Current weather in %s:' % (location))
    print(w[0]['weather'][0]['main'], '-', w[0]['weather'][0]['description'])
    print()
    print('Tomorrow:')
    print(w[1]['weather'][0]['main'], '-', w[1]['weather'][0]['description'])
    print()
    print('Day after tomorrow:')
    print(w[2]['weather'][0]['main'], '-', w[2]['weather'][0]['description'])

def f230():
    # Chapter 15 - Time, Scheduling Tasks, Launching Programs
    print('Chapter 15 - Time, Scheduling Tasks, Launching Programs')

def f231():
    # The time.time() Function

    from time import time
    print(time())
    # number of seconds since 12 am on January 1, 1970 (UTC)
    # The return number is called epoch timestamp

def f232():
    # calculate running time for a program

    import time

    def calcProd():
        # Calculate the product of the first 100,000 numbers.
        product = 1
        for i in range(1, 100000):
            product = product * i
        return product


    startTime = time.time()
    prod = calcProd()
    endTime = time.time()
    print('The result is %s digits long.' % (len(str(prod))))
    print('Took %s seconds to calculate.' % (endTime - startTime))

    # cProfile.run can() be used as an alternative for time.time()
    # For help go to: https://docs.python.org/3/library/profile.html.


def f233():
    # sleep function to pause the program

    import time
    for i in range(3):
        print('Tick')
        time.sleep(1)
        print('Tock')
        time.sleep(1)

     # Making 10 calls to time.sleep(1) is equivalent to
     # time.slweep(10)
    for i in range(10):
        time.sleep(1)

def f234():
    # Rounding the current time number

    import time
    now = time.time()
    print(now)
    print(round(now, 2))
    print(round(now, 4))
    print(round(now)) # No 2nd argument, rounds to nearest integer

def f235():
    # Project: Super Stopwatch - track spent time on a task

    # Step 1: Set Up the Program to Track Times

    import time

    # Display the program's instructions.
    print('''
    Press ENTER to begin. Afterwards,
    press ENTER to "click" the stopwatch.Press Ctrl-C to quit.''')
    input()
    # press Enter to begin
    print('Started.')
    startTime = time.time()
    # get the first lap's start time
    lastTime = startTime
    lapNum = 1

    # Step 2: Track and Print Lap Times

    try:
        # To prevent crashing by Ctl-c, we wrap this part of the -
        # program in a try statement.
        while True:     # infinit loop, to interrupt Ctl-c needed.
            input()
            lapTime = round(time.time() - lastTime, 2)
            totalTime = round(time.time() - startTime, 2)
            print('Lap #%s: Total time: %s - Lap time: %s' % (lapNum, totalTime, lapTime), end='')
            lapNum += 1
            lastTime = time.time() # reset the last lap time

    except KeyboardInterrupt:
        # Handle the Ctrl-C exception to keep its error message from displaying.
        print('\nDone.')

def f236():
    # The datetime Module
    import datetime, time

    # The following returns a datetime object for the current date
    # and time:
    print(type(datetime.datetime.now()))
    print(datetime.datetime.now()) # Displays current date and time

    # If you in IDLE:
    # >>> datetime.datetime.now()
    # datetime.datetime(2019, 5, 6, 17, 2, 33, 739924)

    # To make  datetime object, we use function datetime()
    dt = datetime.datetime(2015, 10, 21, 16, 29, 0)
    print(dt.year, dt.month, dt.day)
    print(dt.hour, dt.minute, dt.second)

    # To convert a Unix epoch timestamp to a datetime object, we -
    # use function fromtimestamp().

    dt = datetime.datetime.fromtimestamp(1000000)
    print('\nPrinting 1,000,000 seconds after the Unix epoch in the'
          + '\nform of a datetime object: ', dt)

    dt = datetime.datetime.fromtimestamp(time.time())
    print('\nPrinting the current time in the form of a'
          + '\ndatetime object: ', dt)

def f237():
    # Compare datetime objects with each other

    import datetime

    halloween2015 = datetime.datetime(2015, 10, 31, 0, 0, 0)
    newyears2016 = datetime.datetime(2016, 1, 1, 0, 0, 0)
    oct31_2015 = datetime.datetime(2015, 10, 31, 0, 0, 0)
    print(halloween2015 == oct31_2015)   # True
    print(halloween2015 > newyears2016)  # False
    print(newyears2016 > halloween2015)  # True
    print(newyears2016 != oct31_2015)    # True

def f238():
    # Find time duration with timedelta module

    import datetime

    # timedelta arguments are: (has no month or year argument)
    # days, hours, minutes, seconds, milliseconds and microseconds
    delta = datetime.timedelta(days=11, hours=10, minutes=9, seconds=8)
    print(delta.days, delta.seconds, delta.microseconds)
    print(delta.total_seconds())
    print(str(delta))             # nice human readable format

    print('.' * 10)

def f239():
    # Using arithmatic operators on datetime values

    import datetime

    # calculate the date 1,000 days from now, using arithmatic op-
    # erators on datetime values

    dt = datetime.datetime.now()
    print(dt)
    thousandDays = datetime.timedelta(days=1000)
    print(dt + thousandDays)

    # timedelta objects can be added or subtracted with datetime -
    # objects or other timedelta objects using the + and - operat-
    # ors. A timedelta object can be multiplied or divided by int-
    # eger or float values with the * and / operators.

    oct21st = datetime.datetime(2015, 10, 21, 16, 29, 0)
    aboutThirtyYears = datetime.timedelta(days=365 * 30)
    print(oct21st)
    print(oct21st - aboutThirtyYears)
    print(oct21st - (2 * aboutThirtyYears))

def f240():
    # Pausing Until a Specific Date

    import datetime
    import time
    halloween2016 = datetime.datetime(2016, 10, 31, 0, 0, 0)
    while datetime.datetime.now() < halloween2016:
        time.sleep(1)

   # The time.sleep(1) call will pause your Python program so that
   # the computer doesn’t waste CPU processing cycles simply chec-
   # king the time over and over. Rather, the while loop will just
   # check the condition once per second and continue with the re-
   # st of the program after Halloween 2016

def f241():
    # Converting datetime Objects into Strings

    import datetime

    # Use strftime() to get a custom string format
    oct21st = datetime.datetime(2015, 10, 21, 16, 29, 0)

    print(oct21st.strftime('%Y/%m/%d %H:%M:%S'))
    print(oct21st.strftime('%I:%M %p'))
    print(oct21st.strftime("%B of '%y"))


    '''
    Directives meaning:
    %Y Year with century, as in '2014'
    %y Year without century, '00' to '99' (1970 to 2069)
    %m Month as a decimal number, '01' to '12'
    %B Full month name, as in 'November'
    %b Abbreviated month name, as in 'Nov'
    %d Day of the month, '01' to '31'
    %j Day of the year, '001' to '366'
    %w Day of the week, '0' (Sunday) to '6' (Saturday)
    %A Full weekday name, as in 'Monday'
    %a Abbreviated weekday name, as in 'Mon'
    %H Hour (24-hour clock), '00' to '23'
    %I Hour (12-hour clock), '01' to '12'
    %M Minute, '00' to '59'
    %S Second, '00' to '59'
    %p 'AM' or 'PM'
    %% Literal '%' character
    '''

def f242():
    # Converting Strings into datetime Objects

    import datetime

    # Parsing strings that are convertible to datetime object can
    # be done using method strptime()

    print('',
    datetime.datetime.strptime('October 21, 2015', '%B %d, %Y'),'\n',
    datetime.datetime.strptime('2015/10/21 16:29:00', '%Y/%m/%d %H:%M:%S'),'\n',
    datetime.datetime.strptime("October of '15", "%B of '%y"),'\n',
    datetime.datetime.strptime("November of '63", "%B of '%y")
    )
    '''
    Directives meaning:
    %Y Year with century, as in '2014'
    %y Year without century, '00' to '99' (1970 to 2069)
    %m Month as a decimal number, '01' to '12'
    %B Full month name, as in 'November'
    %b Abbreviated month name, as in 'Nov'
    %d Day of the month, '01' to '31'
    %j Day of the year, '001' to '366'
    %w Day of the week, '0' (Sunday) to '6' (Saturday)
    %A Full weekday name, as in 'Monday'
    %a Abbreviated weekday name, as in 'Mon'
    %H Hour (24-hour clock), '00' to '23'
    %I Hour (12-hour clock), '01' to '12'
    %M Minute, '00' to '59'
    %S Second, '00' to '59'
    %p 'AM' or 'PM'
    %% Literal '%' character
    '''

def f243():
    # Review of Python’s Time Functions

    print('''
    •    A Unix epoch timestamp (used by the time module) is a fl-
    oat or integer value of the number of seconds  since  12 am on
    January 1, 1970, UTC.

    •    A datetime object ( of the datetime module ) has integers
    stored in the attributes year , month , day , hour ,  minute ,
    and second.

    •    A timedelta object (of the datetime module)  represents a
    time duration, rather than a specific moment.


    Here’s a review of time functions and their parameters and re-
    turn values:
    ..............................................................

    •    The time.time() function returns an epoch timestamp float
    value of the current moment.

    •    The time.sleep(seconds) function   stops  the program for
    the amount of seconds specified by the seconds argument.

    •    The datetime.datetime(year, month, day, hour, minute, se-
    cond) function returns a datetime object of the moment specif-
    ied by the arguments. If hour , minute ,  or  second arguments
    are not provided, they default to 0.

    •    The datetime.datetime.now() function returns a datetime -
    object of the current moment.

    •    The datetime.datetime.fromtimestamp(epoch)  function ret-
    urns a datetime object of the moment represented by  the epoch
    timestamp argument.

    •    The datetime.timedelta(weeks, days, hours, minutes, seco-
    nds, milliseconds, microseconds) function returns  a timedelta
    object representing a duration of time. The function’s keyword
    arguments are all optional and do not include month or year.

    •    The total_seconds() method for  timedelta objects returns
    the number of seconds the timedelta object represents.

    •    The strftime(format) method returns a string  of the time
    represented by the  datetime object in a custom  format that’s
    based on the format string.

    •    The datetime.datetime.strptime(time_string, format) func-
    tion  returns  a datetime object of  the moment  specified  by
    time_string , parsed using the format string argument.
    '''
    )

def f244():
    # Multithreading

    # To make a separate thread, you first need  to make  a Thread
    # object by ­calling the threading.Thread() function.

    # This thread demo shows running different parts  of  a python
    # program at the same time in parallel:

    import threading, time

    print('Start of program.')

    # The function we want to call in the new thread is takeANap()

    def takeANap():
        time.sleep(5)
        print('Wake up!')

    threadObj = threading.Thread(target=takeANap)
    threadObj.start()

    # Notice that the keyword argument is target=takeANap ,
    # not target=takeANap() .

    print('End of program.')

    # output:
    # Start of program.
    # End of program.
    # Wake up!

    # The first is the original thread that began at  the start of
    # the program and ends after print('End of program.') .
    # The 2nd thread is created  when threadObj.start() is called,
    # begins  at  the  start  of the takeANap() function, and ends
    # after takeANap() returns.

def f245():
    # Passing Arguments to the Thread’s Target Function

    import threading


    print('Cats', 'Dogs', 'Frogs', sep=' & ')

    print('#' * 20)

    threadObj = threading.Thread(target=print,
        args=('Cats', 'Dogs', 'Frogs'), kwargs={'sep': ' & '})

    threadObj.start()

    # 'threading' module is higher-level threading interface.
    # 'threading.Thread' is an object that represents a thread.

    # The arguments for Thread object constructer:
    # 'target' is the callable object to be invoked by the run() -
    # method. Defaults to None, meaning nothing is called.
    # 'args' is the argument tuple for the target invocation,
    # defaults to ().
    # 'kwargs' is a dictionary of keyword arguments for the target
    # invocation. Defaults to {}.

    # To avoid concurrency issues, never let multiple threads read
    # or write the same variables.


def f246():
    # Downloads XKCD comics using multiple threads

    import requests, os, bs4, threading

    os.makedirs('xkcd', exist_ok=True)    # store comics in ./xkcd
    def downloadXkcd(startComic, endComic):
        for urlNumber in range(startComic, endComic):
            # Download the page.
            print('Downloading page http://xkcd.com/%s...' % (urlNumber))
            res = requests.get('http://xkcd.com/%s' % (urlNumber))
            res.raise_for_status()

            soup = bs4.BeautifulSoup(res.text)

            # Find the URL of the comic image.
            comicElem = soup.select('#comic img')
            if comicElem == []:
                print('Could not find comic image.')

            else:
                comicUrl = comicElem[0].get('src')
                # Download the image.
                print('Downloading image %s...' % (comicUrl))
                res = requests.get(comicUrl)
                res.raise_for_status()

                # Save the image to ./xkcd.
                imageFile = open(os.path.join('xkcd', os.path.basename(comicUrl)), 'wb')
                for chunk in res.iter_content(100000):
                    imageFile.write(chunk)
                imageFile.close()

    # Create and start the Thread objects.

    downloadThreads = [] # a list of all the Thread objects

    for i in range(0, 1400, 100):     # loops 14 times,
                                      # creates 14 threads

        downloadThread = threading.Thread(target=downloadXkcd,
                                          args=(i, i + 99))
        downloadThreads.append(downloadThread)
        downloadThread.start()


    # Wait for all threads to end.
    for downloadThread in downloadThreads:
        downloadThread.join()
    print('Done.')

def f247():
    # Launching Other Programs from Python
    import subprocess

    subprocess.Popen('/usr/bin/gnome-calculator')

    # Popen object has two methods: poll() and wait()
    # The poll() method will return None  if the process  is still
    # running at the  time poll() is called.  If the  program  has
    # terminated, it will return the process’s integer exit code.

    # The wait() method will block until the launched process has
    # terminated. This is helpful if you want your program to pau-
    # se until the user finishes with the other program. The retu-
    # rn value of wait() is the process’s integer exit code.

    '''
    >>> calcProc = subprocess.Popen('/usr/bin/gnome-calculator')
    >>> calcProc.poll() == None
    True
    >>> calcProc.wait()
    0
    >>> calcProc.poll()  # After closing calculator program
    0
    '''

def f248():
    import subprocess

    # Passing Command Line Arguments to Popen()
    subprocess.Popen(['/usr/bin/gedit', 'list'])

def f249():
    # Opening Files with Default Applications

    fileObj = open('hello.txt', 'w')
    fileObj.write('Hello world!')
    fileObj.close()

    import subprocess
    subprocess.Popen(['start', 'hello.txt'], shell=True)

def f250():
    print('something')

